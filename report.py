import sys
import os
import winreg
import logging
from datetime import datetime
from io import BytesIO

import requests
from requests.auth import HTTPBasicAuth
from bs4 import BeautifulSoup, NavigableString, Tag
from PIL import Image
from docx import Document
from docx.shared import Inches

from PyQt6 import QtWidgets, QtCore, QtGui
from PyQt6.QtCore import Qt, QTimer, QThread, pyqtSignal, QUrl
from PyQt6.QtGui import QIcon, QDesktopServices
from PyQt6.QtWidgets import (
    QWidget, QLabel, QVBoxLayout, QApplication,
    QDialog, QFormLayout, QLineEdit, QComboBox,
    QPushButton, QHBoxLayout, QMessageBox,
    QTextEdit, QFileDialog
)

import pytz
import pyrebase
import firebase_admin
from firebase_admin import credentials, firestore


# ================= LOGGING SETUP =================
LOG_FILENAME = "report_log.txt"
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILENAME, encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ],
)
logger = logging.getLogger("report")

org_url = "https://tfs.alliancewebpos.com/tfs/WebPOSCollection/WebPOS"
collection_url = "https://tfs.alliancewebpos.com/tfs/WebPOSCollection"
template_path = "Report Template.docx"

# ================= AUTO-UPDATE CONFIG =================
APP_VERSION = "3.0.6"
GITHUB_REPO = "shinispades/AutoReport"
GITHUB_API_URL = f"https://api.github.com/repos/{GITHUB_REPO}/releases/latest"


# ================= TOOL REGISTRY =================
# To add a new tool, append a dict to this list:
#   {
#       "id": "unique_tool_id",          # unique identifier
#       "name": "Display Name",          # shown on the card
#       "description": "Short desc...",  # shown on the card
#       "icon": "🔧",                    # emoji or text icon
#       "handler": "method_name",        # method name on LoginWindow to call
#       "category": "General",           # optional category for grouping
#   }
# Then implement the handler method on LoginWindow.
# The Tools page auto-fetches all tools from this list.

TOOL_REGISTRY = [
    # --- Tool 1 ---
    {
        "id": "tool_1",
        "name": "Missing Sales Generator",
        "description": "Generate diagnostic SQL and helper scripts for missing POS sales",
        "icon": "🗃️",
        "handler": "run_tool_1",
        "category": "Reports",
    },
    # --- Tool 2 ---
    {
        "id": "tool_2",
        "name": "Tool 2",
        "description": "Description for Tool 2",
        "icon": "📊",
        "handler": "run_tool_2",
        "category": "General",
    },
    # --- Tool 3 ---
    {
        "id": "tool_3",
        "name": "Tool 3",
        "description": "Description for Tool 3",
        "icon": "📁",
        "handler": "run_tool_3",
        "category": "Reports",
    },
    # --- Tool 4 ---
    {
        "id": "tool_4",
        "name": "Tool 4",
        "description": "Description for Tool 4",
        "icon": "🔍",
        "handler": "run_tool_4",
        "category": "Reports",
    },
    # --- Tool 5 ---
    {
        "id": "tool_5",
        "name": "Tool 5",
        "description": "Description for Tool 5",
        "icon": "⚙️",
        "handler": "run_tool_5",
        "category": "Utilities",
    },
    # --- Tool 6 ---
    {
        "id": "tool_6",
        "name": "Tool 6",
        "description": "Description for Tool 6",
        "icon": "📝",
        "handler": "run_tool_6",
        "category": "Utilities",
    },
    # --- Tool 7 ---
    {
        "id": "tool_7",
        "name": "Tool 7",
        "description": "Description for Tool 7",
        "icon": "🗂️",
        "handler": "run_tool_7",
        "category": "Utilities",
    },
    # --- Tool 8 ---
    {
        "id": "tool_8",
        "name": "Tool 8",
        "description": "Description for Tool 8",
        "icon": "📈",
        "handler": "run_tool_8",
        "category": "Analytics",
    },
    # --- Tool 9 ---
    {
        "id": "tool_9",
        "name": "Tool 9",
        "description": "Description for Tool 9",
        "icon": "🛠️",
        "handler": "run_tool_9",
        "category": "Analytics",
    },
    # --- Tool 10 ---
    {
        "id": "tool_10",
        "name": "Tool 10",
        "description": "Description for Tool 10",
        "icon": "📦",
        "handler": "run_tool_10",
        "category": "Analytics",
    },
    # === ADD MORE TOOLS BELOW ===
    # {
    #     "id": "my_new_tool",
    #     "name": "My New Tool",
    #     "description": "What it does",
    #     "icon": "✨",
    #     "handler": "run_my_new_tool",
    #     "category": "General",
    # },
]


def _fetch_pat_expiry(pat):
    """
    Query the TFS token API to get the expiry date of the given PAT.
    Returns an ISO date string (e.g. '2025-08-01T00:00:00Z') or None.
    """
    url = f"{collection_url}/_apis/tokens/pats?api-version=7.1-preview.1"
    logger.info("Fetching PAT expiry from %s", url)
    try:
        resp = requests.get(url, auth=HTTPBasicAuth("", pat), timeout=15)
        if resp.status_code == 200:
            data = resp.json()
            pats = data.get("patTokens", data.get("value", []))
            logger.debug("Token API returned %d PAT(s)", len(pats))
            # Find the PAT that matches the token we're using (compare first 4 chars)
            short = pat[:4]
            for token in pats:
                display_token = token.get("displayTokenId", "")
                if display_token.startswith(short):
                    expiry = token.get("validTo", "")
                    logger.info("PAT expiry found: %s", expiry)
                    return expiry
            # If no match by prefix, return the first token's expiry
            if pats:
                expiry = pats[0].get("validTo", "")
                logger.info("PAT expiry (first token): %s", expiry)
                return expiry
            logger.warning("No PAT tokens returned from API")
            return None
        elif resp.status_code == 403:
            logger.warning("PAT lacks 'vso.tokens_manage' scope — cannot read expiry (403)")
            return None
        else:
            logger.warning("Token API returned status %d: %s", resp.status_code, resp.text[:200])
            return None
    except Exception:
        logger.exception("Failed to fetch PAT expiry")
        return None


def _check_for_update():
    """
    Check GitHub for the latest release. Returns a dict with
    'version', 'download_url', 'asset_type', and 'body' if an update
    is available, or None if already up-to-date or the check fails.
    """
    logger.info("Checking for update: current=%s, repo=%s", APP_VERSION, GITHUB_REPO)
    try:
        resp = requests.get(GITHUB_API_URL, timeout=10, headers={
            "Accept": "application/vnd.github+json"
        })
        if resp.status_code != 200:
            logger.warning("GitHub API returned %d", resp.status_code)
            return None

        data = resp.json()
        tag = data.get("tag_name", "").lstrip("v")
        if not tag:
            logger.warning("No tag_name in release response")
            return None

        logger.info("Latest release: %s (current: %s)", tag, APP_VERSION)

        # Simple version comparison: split on '.' and compare as ints
        def parse_ver(v):
            try:
                return [int(x) for x in v.split(".")]
            except ValueError:
                return [0]

        if parse_ver(tag) <= parse_ver(APP_VERSION):
            logger.info("Already up-to-date")
            return None

        # Find the first .zip asset (preferred) or .exe asset
        assets = data.get("assets", [])
        download_url = None
        asset_type = None
        for asset in assets:
            name = asset.get("name", "")
            if name.endswith(".zip"):
                download_url = asset.get("browser_download_url")
                asset_type = "zip"
                break
        if not download_url:
            for asset in assets:
                name = asset.get("name", "")
                if name.endswith(".exe"):
                    download_url = asset.get("browser_download_url")
                    asset_type = "exe"
                    break

        if not download_url:
            logger.warning("No .zip or .exe asset found in release %s", tag)
            return None

        logger.info("Update asset found: type=%s, url=%s", asset_type, download_url[:80])
        return {
            "version": tag,
            "download_url": download_url,
            "asset_type": asset_type,
            "body": data.get("body", ""),
        }
    except Exception:
        logger.exception("Update check failed")
        return None


def _download_update(download_url, progress_callback=None):
    """
    Download the update file to a temp location. Returns the file path or None.
    progress_callback(bytes_downloaded, total_bytes) is called periodically.
    """
    import tempfile
    logger.info("Downloading update from %s", download_url)
    try:
        resp = requests.get(download_url, stream=True, timeout=60)
        resp.raise_for_status()
        total = int(resp.headers.get("content-length", 0))
        downloaded = 0

        tmp_dir = tempfile.gettempdir()
        filename = download_url.split("/")[-1].split("?")[0]
        tmp_path = os.path.join(tmp_dir, filename)

        with open(tmp_path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)
                downloaded += len(chunk)
                if progress_callback and total:
                    progress_callback(downloaded, total)

        logger.info("Update downloaded to %s (%d bytes)", tmp_path, downloaded)
        return tmp_path
    except Exception:
        logger.exception("Failed to download update")
        return None


def _apply_zip_update(zip_path):
    """
    Extract a zip update and replace files in the current app directory.
    Locked files (e.g. the running exe) are staged as .new and swapped
    by a batch script after the app exits.
    Returns True on success, False on failure.
    """
    import zipfile
    import shutil

    if getattr(sys, 'frozen', False):
        app_dir = os.path.dirname(sys.executable)
        exe_name = os.path.basename(sys.executable)
    else:
        app_dir = os.path.dirname(os.path.abspath(__file__))
        exe_name = None

    logger.info("Applying zip update: %s -> %s", zip_path, app_dir)

    try:
        extract_dir = os.path.join(os.path.dirname(zip_path), "update_extract")
        if os.path.exists(extract_dir):
            shutil.rmtree(extract_dir)
        os.makedirs(extract_dir)

        with zipfile.ZipFile(zip_path, 'r') as zf:
            zf.extractall(extract_dir)
            logger.info("Extracted %d files from zip", len(zf.namelist()))

        # Copy extracted files over the app directory
        copied = 0
        locked_files = []
        for item in os.listdir(extract_dir):
            src = os.path.join(extract_dir, item)
            dst = os.path.join(app_dir, item)
            try:
                if os.path.isdir(src):
                    if os.path.exists(dst):
                        shutil.rmtree(dst)
                    shutil.copytree(src, dst)
                else:
                    shutil.copy2(src, dst)
                copied += 1
                logger.debug("Replaced: %s", dst)
            except PermissionError:
                # File is locked (likely the running exe) — stage as .new
                staged = dst + ".new"
                try:
                    # Remove stale .new file from a previous failed update
                    if os.path.exists(staged):
                        os.remove(staged)
                except OSError:
                    logger.warning("Could not remove stale .new file: %s", staged)
                try:
                    shutil.copy2(src, staged)
                    locked_files.append((staged, dst))
                    copied += 1
                    logger.info("Staged locked file: %s -> %s", src, staged)
                except Exception:
                    logger.exception("Failed to stage locked file: %s -> %s", src, staged)

        # If there are locked files, create a batch script to swap them after exit
        if locked_files:
            bat_path = os.path.join(app_dir, "_update_swap.bat")
            pid = os.getpid()
            lines = [
                "@echo off",
                f"taskkill /PID {pid} /F >nul 2>&1",
                "timeout /t 2 /nobreak >nul",
            ]
            for staged, dst in locked_files:
                lines.append(f'del /f /q "{dst}" 2>nul')
                lines.append(f'rename "{staged}" "{os.path.basename(dst)}"')
            lines.append(f'del /f /q "{bat_path}" 2>nul')
            lines.append(f'start "" "{os.path.join(app_dir, exe_name or "report.py")}"')
            with open(bat_path, "w") as f:
                f.write("\n".join(lines))
            logger.info("Created update swap script: %s", bat_path)

        # Cleanup
        shutil.rmtree(extract_dir, ignore_errors=True)
        os.remove(zip_path)
        logger.info("Zip update applied: %d item(s) replaced", copied)
        return True

    except Exception:
        logger.exception("Failed to apply zip update")
        return False


# Firebase — lazy initialized on first use so the UI appears instantly
_firebase_config = {
  "apiKey": "AIzaSyD-uAl1Ce-2NRoAVVYExH5-MSJ6K4y-LEo",
  "authDomain": "autoreport-db4a7.firebaseapp.com",
  "projectId": "autoreport-db4a7",
  "databaseURL": "",
  "storageBucket": "autoreport-db4a7.firebasestorage.app",
  "messagingSenderId": "173593376987",
  "appId": "1:173593376987:web:66412a51ebb20818a9b4fe",
  "measurementId": "G-F2VL1BBRGT"
}
_auth = None
_db = None

def _get_firebase():
    """Initialize Firebase on first call, then cache the result."""
    global _auth, _db
    if _auth is None:
        logger.info("Initializing Firebase (first call)")
        try:
            firebase = pyrebase.initialize_app(_firebase_config)
            _auth = firebase.auth()
            cred = credentials.Certificate("serviceAccountKey.json")
            firebase_admin.initialize_app(cred)
            _db = firestore.client()
            logger.info("Firebase initialized successfully")
        except Exception:
            logger.exception("Firebase initialization failed")
            raise
    return _auth, _db



def replace_placeholders_in_paragraph(paragraph, replacements: dict):
    """
    Robustly replace {{PLACEHOLDER}} tokens in a paragraph, even when
    Word has split the token across multiple runs.
    """
    full_text = paragraph.text
    if not any(ph in full_text for ph in replacements):
        return  # nothing to do, skip untouched paragraphs (preserves formatting)

    logger.debug("Paragraph contains placeholder(s): %r", full_text[:120])

    # Apply all replacements to the combined text
    new_text = full_text
    for ph, value in replacements.items():
        new_text = new_text.replace(ph, value)

    if new_text == full_text:
        return

    # Wipe all runs except the first, then set the first run's text.
    # This collapses formatting to whatever the first run had, which is
    # usually fine for placeholder tokens (they're typically uniformly formatted).
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def replace_placeholders_in_doc(doc, replacements: dict):
    logger.info("Replacing placeholders in document (%d replacements)", len(replacements))
    logger.debug("Placeholder keys: %s", list(replacements.keys()))

    # Paragraphs outside tables
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, replacements)

    # Paragraphs inside tables (including nested tables)
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, replacements)
                # handle nested tables inside a cell
                for nested_table in cell.tables:
                    for nrow in nested_table.rows:
                        for ncell in nrow.cells:
                            for npara in ncell.paragraphs:
                                replace_placeholders_in_paragraph(npara, replacements)

    # --- Text boxes (w:txbxContent) ---
    # python-docx does NOT expose paragraphs inside text boxes via
    # doc.paragraphs/doc.tables, since they live inside drawing/VML
    # shapes (mc:AlternateContent -> mc:Choice/w:drawing and
    # mc:Fallback -> w:pict). Both copies must be patched identically,
    # so we walk the raw XML tree directly.
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    txbx_count = 0
    for txbx_content in body.iter(qn('w:txbxContent')):
        txbx_count += 1
        for p in txbx_content.iter(qn('w:p')):
            wrapped_para = Paragraph(p, None)
            replace_placeholders_in_paragraph(wrapped_para, replacements)
    logger.debug("Processed %d text box(es) in document", txbx_count)


def add_html_content_to_paragraph(paragraph, html_content, pat=""):
    """Render HTML content (with images) into a Word paragraph."""
    logger.debug("Rendering HTML content into paragraph (length=%d chars)", len(html_content))
    soup = BeautifulSoup(html_content, "html.parser")
    img_count = 0
    img_fail_count = 0

    def recurse(node, para):
        nonlocal img_count, img_fail_count
        if isinstance(node, NavigableString):
            for line in str(node).splitlines():
                para.add_run(line)
                para.add_run().add_break()
        elif isinstance(node, Tag):
            if node.name == "br":
                para.add_run().add_break()
            elif node.name == "p":
                for child in node.children:
                    recurse(child, para)
                para.add_run().add_break()
            elif node.name == "li":
                para.add_run("• ")
                for child in node.children:
                    recurse(child, para)
                para.add_run().add_break()
            elif node.name == "img":
                img_url = node.get("src")
                if img_url:
                    try:
                        logger.debug("Fetching image: %s", img_url[:120])
                        img_resp = requests.get(img_url, auth=HTTPBasicAuth("", pat))
                        img_resp.raise_for_status()
                        image_stream = BytesIO(img_resp.content)
                        image = Image.open(image_stream)
                        dpi = image.info.get("dpi", (96, 96))[0]
                        width_inches = image.width / dpi
                        max_width_inches = 6
                        image_stream.seek(0)
                        if width_inches > max_width_inches:
                            para.add_run().add_picture(image_stream, width=Inches(max_width_inches))
                        else:
                            para.add_run().add_picture(image_stream)
                        img_count += 1
                        logger.debug("Image inserted successfully (dpi=%s, width=%.2f in)", dpi, width_inches)
                    except Exception as e:
                        img_fail_count += 1
                        logger.error("Error inserting image from %s: %s", img_url[:120], e)
            else:
                for child in node.children:
                    recurse(child, para)

    for child in soup.children:
        recurse(child, paragraph)

    logger.info("HTML rendering complete: %d image(s) inserted, %d failed", img_count, img_fail_count)


# ================= CUSTOM DIALOG =================
class AppDialog(QtWidgets.QDialog):
    def __init__(self, title, message, parent=None):
        super().__init__(parent)

        self.setWindowIcon(QtGui.QIcon("asset/death.ico"))
        self.setWindowTitle(title)
        self.setFixedSize(320, 160)
        self.setObjectName("Root")

        # Use the same dark/light styles as your main window
        if is_windows_dark_mode():
            self.setStyleSheet("""
            QDialog#Root {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1,
                    stop:0 #07101a, stop:1 #07121a);
                font-family: Inter, Segoe UI, Arial;
                color: #e6eef5;
            }
            QLabel {
                color: #e6eef5;
                background: transparent;
            }
            QPushButton {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1,
                    stop:0 #ed3b3b, stop:1 #c62e2e);
                color: white;
                border-radius: 10px;
                padding: 8px 16px;
            }
            QPushButton:hover {
                opacity: 0.92;
            }
            """)
        else:
            self.setStyleSheet("""
            QDialog#Root {
                background: #f5f7fa;
                font-family: Inter, Segoe UI, Arial;
                color: #1f2933;
            }
            QLabel {
                color: #111827;
                background: transparent;
            }
            QPushButton {
                background: #ed3b3b;
                color: white;
                border-radius: 10px;
                padding: 8px 16px;
            }
            QPushButton:hover {
                background: #c62e2e;
            }
            """)
        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(14)
        label = QtWidgets.QLabel(message)
        label.setWordWrap(True)
        button = QtWidgets.QPushButton("OK")
        button.setFixedHeight(34)
        button.clicked.connect(self.accept)
        layout.addWidget(label)
        layout.addStretch()
        layout.addWidget(button)


class UpdateDialog(QtWidgets.QDialog):
    """Dialog shown when an update is available."""
    def __init__(self, version, release_notes, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Update Available")
        self.setFixedSize(400, 280)
        self.setObjectName("UpdateDialog")
        self.download_url = None

        dark = is_windows_dark_mode()
        if dark:
            self.setStyleSheet("""
            QDialog#UpdateDialog {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1,
                    stop:0 #07101a, stop:1 #07121a);
                font-family: Inter, Segoe UI, Arial;
                color: #e6eef5;
            }
            QLabel { color: #e6eef5; background: transparent; }
            QPushButton {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1,
                    stop:0 #ed3b3b, stop:1 #c62e2e);
                color: white; border-radius: 10px; padding: 8px 16px;
            }
            QPushButton:hover { background: #c62e2e; }
            QPushButton#SkipBtn {
                background: transparent; border: 1px solid rgba(255,255,255,0.15);
                color: #98a0a6;
            }
            QPushButton#SkipBtn:hover { border: 1px solid #98a0a6; color: #e6eef5; }
            QProgressBar {
                background: rgba(255,255,255,0.05); border: 1px solid rgba(255,255,255,0.1);
                border-radius: 6px; text-align: center; color: #e6eef5; height: 20px;
            }
            QProgressBar::chunk { background: #ed3b3b; border-radius: 5px; }
            """)
        else:
            self.setStyleSheet("""
            QDialog#UpdateDialog {
                background: #f5f7fa;
                font-family: Inter, Segoe UI, Arial;
                color: #1f2933;
            }
            QLabel { color: #111827; background: transparent; }
            QPushButton {
                background: #ed3b3b; color: white; border-radius: 10px; padding: 8px 16px;
            }
            QPushButton:hover { background: #c62e2e; }
            QPushButton#SkipBtn {
                background: transparent; border: 1px solid #d1d5db; color: #6b7280;
            }
            QPushButton#SkipBtn:hover { border: 1px solid #6b7280; color: #111827; }
            QProgressBar {
                background: #e5e7eb; border: 1px solid #d1d5db;
                border-radius: 6px; text-align: center; color: #111827; height: 20px;
            }
            QProgressBar::chunk { background: #ed3b3b; border-radius: 5px; }
            """)

        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(12)

        title = QtWidgets.QLabel(f"Version {version} is available")
        title.setStyleSheet("font-size: 18px; font-weight: 700;")
        layout.addWidget(title)

        current_label = QtWidgets.QLabel(f"Current version: {APP_VERSION}")
        current_label.setStyleSheet("font-size: 12px; color: #6b7280;")
        layout.addWidget(current_label)

        if release_notes:
            notes_label = QtWidgets.QLabel(release_notes[:200])
            notes_label.setWordWrap(True)
            notes_label.setStyleSheet("font-size: 12px; color: #98a0a6;")
            layout.addWidget(notes_label)

        layout.addSpacing(8)

        self.progress_bar = QtWidgets.QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        self.status_label = QtWidgets.QLabel("")
        self.status_label.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.status_label.setVisible(False)
        layout.addWidget(self.status_label)

        layout.addStretch()

        btn_layout = QtWidgets.QHBoxLayout()
        self.download_btn = QtWidgets.QPushButton("Download Update")
        self.download_btn.setFixedHeight(38)
        self.download_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        self.download_btn.clicked.connect(self._on_download)

        skip_btn = QtWidgets.QPushButton("Skip")
        skip_btn.setObjectName("SkipBtn")
        skip_btn.setFixedHeight(38)
        skip_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        skip_btn.clicked.connect(self.reject)

        btn_layout.addWidget(skip_btn)
        btn_layout.addStretch()
        btn_layout.addWidget(self.download_btn)
        layout.addLayout(btn_layout)

    def _on_download(self):
        self.download_btn.setEnabled(False)
        self.download_btn.setText("Downloading...")
        self.progress_bar.setVisible(True)
        self.status_label.setVisible(True)
        self.status_label.setText("Starting download...")
        QtWidgets.QApplication.processEvents()

        def on_progress(done, total):
            pct = int(done / total * 100)
            self.progress_bar.setValue(pct)
            mb_done = done / (1024 * 1024)
            mb_total = total / (1024 * 1024)
            self.status_label.setText(f"{mb_done:.1f} MB / {mb_total:.1f} MB")
            QtWidgets.QApplication.processEvents()

        # Get update info from the parent
        update_info = getattr(self.parent(), "_update_info", None) if self.parent() else None
        if not update_info:
            self.status_label.setText("Error: No update info")
            self.download_btn.setEnabled(True)
            self.download_btn.setText("Download Update")
            return

        url = update_info["download_url"]
        asset_type = update_info.get("asset_type", "exe")

        path = _download_update(url, progress_callback=on_progress)
        if not path:
            self.status_label.setText("Download failed. Check logs.")
            self.download_btn.setEnabled(True)
            self.download_btn.setText("Retry")
            self.download_btn.clicked.disconnect()
            self.download_btn.clicked.connect(self._on_download)
            return

        self.progress_bar.setValue(100)

        if asset_type == "zip":
            # Extract and replace files in-place
            self.status_label.setText("Extracting update...")
            self.download_btn.setText("Applying update...")
            QtWidgets.QApplication.processEvents()

            if _apply_zip_update(path):
                self.status_label.setText("Update applied! Restarting app...")
                QtWidgets.QApplication.processEvents()
                logger.info("Zip update applied, restarting app")
                self._restart_app()
            else:
                self.status_label.setText("Failed to apply update. Check logs.")
                self.download_btn.setText("Close")
                self.download_btn.setEnabled(True)
                self.download_btn.clicked.connect(self.accept)
        else:
            # .exe — launch installer and exit
            self.status_label.setText("Download complete! Launching installer...")
            QtWidgets.QApplication.processEvents()
            logger.info("Launching update: %s", path)
            try:
                import subprocess
                subprocess.Popen([path], shell=False)
                logger.info("Update installer launched, exiting app")
                QtWidgets.QApplication.quit()
            except Exception:
                logger.exception("Failed to launch update installer")
                self.status_label.setText(f"Downloaded to: {path}")
                self.download_btn.setText("Close")
                self.download_btn.setEnabled(True)
                self.download_btn.clicked.connect(self.accept)

    def _restart_app(self):
        """Restart the application after a zip update."""
        import subprocess

        if getattr(sys, 'frozen', False):
            app_dir = os.path.dirname(sys.executable)
        else:
            app_dir = os.path.dirname(os.path.abspath(__file__))

        bat_path = os.path.join(app_dir, "_update_swap.bat")

        # If a swap script exists (locked files), just quit — the bat handles restart
        if os.path.isfile(bat_path):
            logger.info("Swap script found, launching: %s", bat_path)
            subprocess.Popen(["cmd", "/c", bat_path], shell=False,
                             creationflags=subprocess.CREATE_NO_WINDOW)
            QtWidgets.QApplication.quit()
            return

        # No locked files — restart directly
        exe = sys.executable
        logger.info("Restarting: %s %s", exe, sys.argv)
        try:
            subprocess.Popen([exe] + sys.argv)
            QtWidgets.QApplication.quit()
        except Exception:
            logger.exception("Failed to restart app")
            self.status_label.setText("Update applied. Please restart manually.")
            self.download_btn.setText("Close")
            self.download_btn.setEnabled(True)
            self.download_btn.clicked.connect(self.accept)


class PBIFormDialog(QDialog):
        def __init__(self, parent=None):
            super().__init__(parent)
            self.setWindowTitle("Create Azure DevOps PBI")
            self.setFixedWidth(380)

            if is_windows_dark_mode():
                self.setStyleSheet("""
                QDialog { background: #0f1923; color: #e6eef5; font-family: Inter, Segoe UI, Arial; }
                QLabel { color: #e6eef5; }
                QLineEdit {
                    background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.06);
                    border-radius: 10px; padding: 12px; color: #e6eef5; font-size: 14px;
                }
                QLineEdit:focus { border: 1px solid #ed3b3b; }
                QComboBox {
                    background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.06);
                    border-radius: 10px; padding: 12px 36px 12px 12px; color: #e6eef5;
                    font-size: 14px; min-height: 20px;
                }
                QComboBox:focus { border: 1px solid #ed3b3b; }
                QComboBox::drop-down {
                    subcontrol-origin: padding; subcontrol-position: center right;
                    width: 32px; border: none;
                    border-left: 1px solid rgba(255,255,255,0.06);
                    border-top-right-radius: 10px; border-bottom-right-radius: 10px;
                }
                QComboBox::down-arrow { width: 10px; height: 10px; }
                QComboBox QAbstractItemView {
                    background: #15212c; border: 1px solid #2a3a48; border-radius: 8px;
                    padding: 4px; color: #e6eef5; selection-background-color: #1e3a50;
                    selection-color: #ffffff; outline: none;
                }
                QComboBox QAbstractItemView::item {
                    padding: 8px 12px; border-radius: 4px; min-height: 24px;
                }
                QComboBox QAbstractItemView::item:hover { background: #1e3a50; color: #ffffff; }
                QPushButton {
                    background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #ed3b3b, stop:1 #c62e2e);
                    color: white; border-radius: 10px; padding: 12px 24px; font-size: 14px; font-weight: 600;
                }
                QPushButton:hover { background: #c62e2e; }
                """)
            else:
                self.setStyleSheet("""
                QDialog { background: #f5f7fa; color: #1f2933; font-family: Inter, Segoe UI, Arial; }
                QLabel { color: #111827; }
                QLineEdit {
                    background: white; border: 1px solid #d1d5db;
                    border-radius: 10px; padding: 12px; color: #111827; font-size: 14px;
                }
                QLineEdit:focus { border: 1px solid #ed3b3b; }
                QComboBox {
                    background: white; border: 1px solid #d1d5db;
                    border-radius: 10px; padding: 12px 36px 12px 12px; color: #111827;
                    font-size: 14px; min-height: 20px;
                }
                QComboBox:focus { border: 1px solid #ed3b3b; }
                QComboBox::drop-down {
                    subcontrol-origin: padding; subcontrol-position: center right;
                    width: 32px; border: none;
                    border-left: 1px solid #d1d5db;
                    border-top-right-radius: 10px; border-bottom-right-radius: 10px;
                }
                QComboBox::down-arrow { width: 10px; height: 10px; }
                QComboBox QAbstractItemView {
                    background: white; border: 1px solid #d1d5db; border-radius: 8px;
                    padding: 4px; color: #111827; selection-background-color: #f0f0f0;
                    selection-color: #111827; outline: none;
                }
                QComboBox QAbstractItemView::item {
                    padding: 8px 12px; border-radius: 4px; min-height: 24px;
                }
                QComboBox QAbstractItemView::item:hover { background: #f3f4f6; color: #111827; }
                QPushButton {
                    background: #ed3b3b; color: white; border-radius: 10px;
                    padding: 12px 24px; font-size: 14px; font-weight: 600;
                }
                QPushButton:hover { background: #c62e2e; }
                """)

            layout = QVBoxLayout(self)
            form = QFormLayout()

            # === Back Job (Yes / No) ===
            self.back_job = QComboBox()
            self.back_job.addItems(["Yes", "No"])

            # === Billing Status ===
            self.billing_status = QComboBox()
            self.billing_status.addItems([
                "APS",
                "Billed",
                "Pending From Previous Task",
                "Warranty"
            ])

            # === Product Type ===
            self.product_type = QComboBox()
            self.product_type.addItems([
                "CloudPOS",
                "HRMS",
                "Portfolio ERP/PY",
                "POS V3/V5",
                "SAP",
                "WebPOS"
            ])

            # === Support Type ===
            self.support_type = QComboBox()
            self.support_type.addItems([
                "Onsite",
                "Remote/Offsite"
            ])

            # === Confirmed By ===
            self.confirmed_by = QLineEdit()

            # Add fields to form
            form.addRow("Back Job:", self.back_job)
            form.addRow("Billing Status:", self.billing_status)
            form.addRow("Product Type:", self.product_type)
            form.addRow("Support Type:", self.support_type)
            form.addRow("Confirmed By:", self.confirmed_by)

            layout.addLayout(form)

            # === Buttons ===
            # btn_layout = QHBoxLayout()
            # self.ok_btn = QPushButton("Create PBI")
            # cancel_btn = QPushButton("Cancel")

            # self.ok_btn.clicked.connect(self.validate_and_accept)
            # cancel_btn.clicked.connect(self.reject)

            # btn_layout.addWidget(self.ok_btn)
            # btn_layout.addWidget(cancel_btn)
            # layout.addLayout(btn_layout)

            btn_layout = QHBoxLayout()

            self.ok_btn = QPushButton("Create PBI")
            cancel_btn = QPushButton("Cancel")

            self.ok_btn.clicked.connect(self.validate_and_accept)
            cancel_btn.clicked.connect(self.reject)

            # Keep buttons compact
            self.ok_btn.setFixedWidth(120)
            cancel_btn.setFixedWidth(90)

            btn_layout.addStretch()
            btn_layout.addWidget(self.ok_btn)
            btn_layout.addSpacing(10)
            btn_layout.addWidget(cancel_btn)
            btn_layout.addStretch()

            layout.addLayout(btn_layout)


        def validate_and_accept(self):
            """Basic validation like CLI version"""
            if not self.confirmed_by.text().strip():
                AppDialog(
                    "Validation Error",
                    "Confirmed By cannot be empty.",
                    self
                ).exec()
                return
            self.accept()

        def get_values(self):
            return {
                "back_job": self.back_job.currentText(),   # Yes / No (string, same as CLI)
                "billing_status": self.billing_status.currentText(),
                "product_type": self.product_type.currentText(),
                "support_type": self.support_type.currentText(),
                "confirmed_by": self.confirmed_by.text().strip()
            }



class PBICreatedDialog(QDialog):
    def __init__(self, pbi_id, pbi_url, parent=None):
        super().__init__(parent)

        self.setWindowTitle("PBI Created Successfully")
        self.setFixedWidth(420)

        if is_windows_dark_mode():
            self.setStyleSheet("""
            QDialog { background: #0f1923; color: #e6eef5; font-family: Inter, Segoe UI, Arial; }
            QLabel { color: #e6eef5; }
            QLabel a { color: #60a5fa; }
            QPushButton {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #ed3b3b, stop:1 #c62e2e);
                color: white; border-radius: 10px; padding: 8px 16px; font-size: 14px; font-weight: 600;
            }
            QPushButton:hover { background: #c62e2e; }
            """)
        else:
            self.setStyleSheet("""
            QDialog { background: #f5f7fa; color: #1f2933; font-family: Inter, Segoe UI, Arial; }
            QLabel { color: #111827; }
            QLabel a { color: #2563eb; }
            QPushButton {
                background: #ed3b3b; color: white; border-radius: 10px;
                padding: 8px 16px; font-size: 14px; font-weight: 600;
            }
            QPushButton:hover { background: #c62e2e; }
            """)

        layout = QVBoxLayout(self)
        form = QFormLayout()

        pbi_id_label = QLabel(str(pbi_id))
        pbi_id_label.setTextInteractionFlags(
            Qt.TextInteractionFlag.TextSelectableByMouse
        )

        pbi_url_label = QLabel(f'<a href="{pbi_url}">{pbi_url}</a>')
        pbi_url_label.setOpenExternalLinks(True)
        pbi_url_label.setTextInteractionFlags(
            Qt.TextInteractionFlag.TextSelectableByMouse
        )

        form.addRow("PBI ID:", pbi_id_label)
        form.addRow("PBI URL:", pbi_url_label)

        layout.addLayout(form)

        btn_layout = QHBoxLayout()
        open_btn = QPushButton("Open in Browser")
        close_btn = QPushButton("Close")

        open_btn.clicked.connect(
            lambda: QDesktopServices.openUrl(QUrl(pbi_url))
        )
        close_btn.clicked.connect(self.accept)

        btn_layout.addStretch()
        btn_layout.addWidget(open_btn)
        btn_layout.addWidget(close_btn)

        layout.addLayout(btn_layout)


class MissingSalesDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Missing Sales Generator")
        self.setMinimumWidth(560)

        dark = is_windows_dark_mode()
        if dark:
            self.setStyleSheet("""
            QDialog { background: #0f1923; color: #e6eef5; font-family: Inter, Segoe UI, Arial; }
            QLabel { color: #e6eef5; background: transparent; }
            QLineEdit {
                background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.06);
                border-radius: 8px; padding: 10px 12px; color: #e6eef5; font-size: 13px;
            }
            QLineEdit:focus { border: 1px solid #ed3b3b; }
            QPushButton {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #ed3b3b, stop:1 #c62e2e);
                color: white; border-radius: 10px; padding: 10px 20px; font-size: 14px; font-weight: 600;
            }
            QPushButton:hover { background: #c62e2e; }
            QPushButton#GhostBtn {
                background: transparent; border: 1px solid rgba(255,255,255,0.15);
                color: #98a0a6;
            }
            QPushButton#GhostBtn:hover { border: 1px solid #98a0a6; color: #e6eef5; }
            QTextEdit {
                background: #111a22; border: 1px solid #2a3a48; border-radius: 8px;
                color: #e6eef5; font-size: 12px; font-family: Consolas, monospace;
            }
            """)
        else:
            self.setStyleSheet("""
            QDialog { background: #f5f7fa; color: #1f2933; font-family: Inter, Segoe UI, Arial; }
            QLabel { color: #111827; background: transparent; }
            QLineEdit {
                background: white; border: 1px solid #d1d5db;
                border-radius: 8px; padding: 10px 12px; color: #111827; font-size: 13px;
            }
            QLineEdit:focus { border: 1px solid #ed3b3b; }
            QPushButton {
                background: #ed3b3b; color: white; border-radius: 10px;
                padding: 10px 20px; font-size: 14px; font-weight: 600;
            }
            QPushButton:hover { background: #c62e2e; }
            QPushButton#GhostBtn {
                background: transparent; border: 1px solid #d1d5db; color: #6b7280;
            }
            QPushButton#GhostBtn:hover { border: 1px solid #6b7280; color: #111827; }
            QTextEdit {
                background: white; border: 1px solid #d1d5db; border-radius: 8px;
                color: #111827; font-size: 12px; font-family: Consolas, monospace;
            }
            """)

        tc = parent._theme_colors() if parent and hasattr(parent, "_theme_colors") else {}

        layout = QVBoxLayout(self)
        layout.setContentsMargins(28, 24, 28, 24)
        layout.setSpacing(0)

        # --- Header ---
        header = QLabel("Missing Sales Generator")
        header.setStyleSheet("font-size: 18px; font-weight: 700;")
        layout.addWidget(header)
        layout.addSpacing(4)

        subtitle = QLabel("Generate diagnostic SQL and helper scripts for missing POS sales")
        subtitle.setStyleSheet("color: #6b7280; font-size: 12px;")
        subtitle.setWordWrap(True)
        layout.addWidget(subtitle)

        layout.addSpacing(20)

        # --- Divider ---
        divider = QtWidgets.QFrame()
        divider.setFixedHeight(1)
        divider.setStyleSheet("background: rgba(255,255,255,0.08) if dark else #e5e7eb; border: none;")
        divider.setStyleSheet(f"background: {'rgba(255,255,255,0.08)' if dark else '#e5e7eb'}; border: none;")
        layout.addWidget(divider)
        layout.addSpacing(16)

        # --- Section: Terminal Info ---
        section1 = QLabel("TERMINAL INFORMATION")
        section1.setStyleSheet("color: #ed3b3b; font-size: 11px; font-weight: 700; letter-spacing: 1.5px;")
        layout.addWidget(section1)
        layout.addSpacing(12)

        field_label_style = "color: #98a0a6; font-size: 12px; font-weight: 600;" if dark else "color: #374151; font-size: 12px; font-weight: 600;"

        def make_field(label_text, widget):
            block = QVBoxLayout()
            block.setSpacing(5)
            lbl = QLabel(label_text)
            lbl.setStyleSheet(field_label_style)
            block.addWidget(lbl)
            block.addWidget(widget)
            return block

        self.company_id = QLineEdit()
        self.company_id.setPlaceholderText("WPOS-26012789")
        self.terminal = QLineEdit()
        self.terminal.setPlaceholderText("1 or 0001")
        self.year = QLineEdit()
        self.year.setPlaceholderText("YYYY")
        self.month = QLineEdit()
        self.month.setPlaceholderText("MM")

        grid1 = QtWidgets.QGridLayout()
        grid1.setHorizontalSpacing(16)
        grid1.setVerticalSpacing(12)
        grid1.addLayout(make_field("Company ID", self.company_id), 0, 0)
        grid1.addLayout(make_field("Terminal", self.terminal), 0, 1)
        grid1.addLayout(make_field("Missing Year", self.year), 1, 0)
        grid1.addLayout(make_field("Missing Month", self.month), 1, 1)
        layout.addLayout(grid1)

        layout.addSpacing(20)

        # --- Divider 2 ---
        divider2 = QtWidgets.QFrame()
        divider2.setFixedHeight(1)
        divider2.setStyleSheet(f"background: {'rgba(255,255,255,0.08)' if dark else '#e5e7eb'}; border: none;")
        layout.addWidget(divider2)
        layout.addSpacing(16)

        # --- Section: Connection ---
        section2 = QLabel("CONNECTION")
        section2.setStyleSheet("color: #ed3b3b; font-size: 11px; font-weight: 700; letter-spacing: 1.5px;")
        layout.addWidget(section2)
        layout.addSpacing(12)

        self.link = QLineEdit()
        self.link.setPlaceholderText("https://example.com")
        layout.addLayout(make_field("Company Link", self.link))

        layout.addSpacing(24)

        # --- Buttons ---
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)

        copy_btn = QPushButton("\U0001f4cb Copy")
        copy_btn.setObjectName("GhostBtn")
        copy_btn.setFixedHeight(40)
        copy_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        copy_btn.clicked.connect(self.copy_to_clipboard)

        self.generate_btn = QPushButton("Generate & Save")
        self.generate_btn.setFixedHeight(40)
        self.generate_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        self.generate_btn.clicked.connect(self.generate_and_save)

        btn_layout.addStretch()
        btn_layout.addWidget(copy_btn)
        btn_layout.addWidget(self.generate_btn)
        layout.addLayout(btn_layout)

        layout.addSpacing(16)

        # --- Preview ---
        preview_label = QLabel("PREVIEW")
        preview_label.setStyleSheet("color: #ed3b3b; font-size: 11px; font-weight: 700; letter-spacing: 1.5px;")
        layout.addWidget(preview_label)
        layout.addSpacing(8)

        self.preview = QTextEdit()
        self.preview.setReadOnly(True)
        self.preview.setMinimumHeight(100)
        self.preview.setMaximumHeight(160)
        self.preview.setPlaceholderText("Generated script will appear here...")
        layout.addWidget(self.preview)

    def _pad_terminal(self, terminal):
        t = terminal.strip()
        if len(t) == 1:
            return "000" + t
        elif len(t) == 2:
            return "00" + t
        elif len(t) == 3:
            return "0" + t
        return t

    def _build_content(self):
        cid = self.company_id.text().strip()
        terminal = self._pad_terminal(self.terminal.text())
        year = self.year.text().strip()
        month = self.month.text().strip()
        link = self.link.text().strip()
        missing_date = year + month + "01"
        missing_date2 = year + "1231"

        lacking = (
            'WITH LACKING:\n'
            'SELECT fzcounter, fsale_date, MIN(fdocument_no) AS fmin, MAX(fdocument_no) AS fmax, '
            'SUM(fgross) AS fgross, COUNT(fdocument_no) AS fdocument_no, '
            'MAX(fdocument_no) - MIN(fdocument_no) + 1 AS Expected_OR_count, '
            'count(*)-(max(fdocument_no)-min(fdocument_no)+1) as missing_OR '
            'FROM pos_sale WHERE fcompanyid="' + cid + '" and ftermid="' + terminal + '" '
            'AND fsale_date>="' + missing_date + '" AND fdocument_no <> 0 '
            'GROUP BY fsale_date, fzcounter'
        )

        full_missing = (
            '\n\nFULL MISSING SCRIPT:\n'
            'select fsale_date, fzcounter, min(fdocument_no) as min_fdoc, max(fdocument_no) as max_fdoc, '
            'count(*) as trxcnt,(max(fdocument_no) - min(fdocument_no) + 1) as expected, '
            "case when count(*) = (max(fdocument_no) - min(fdocument_no) + 1) then 'Y' else 'N' end as same, "
            'count(*) - (max(fdocument_no) - min(fdocument_no) + 1) as lacking, sum(fgross) as fgross '
            'from pos_sale where fcompanyid="' + cid + '" AND ftermid="' + terminal + '" '
            "and fsale_date >= '" + missing_date + "' and fsale_date <= '" + missing_date2 + "' "
            "and fdocument_no <> '0' group by fsale_date, fzcounter having (same = 'N')"
        )

        excel = (
            '\n\nFOR EXCEL:\n'
            '=IF(E8-E7=1,"OK","MISSING")\n'
            '=IF(E8 - E7 > 1, IF(E8-E7 > 2, E7 + 1 & "-" & E8 - 1, E8 -1), "OK")'
        )

        today = datetime.now().strftime("%Y%m%d")
        select_or = (
            '\n\nSELECT OR/DATE SCRIPT:\n'
            'select * from pos_sale where fcompanyid="' + cid + '" and ftermid="' + terminal + '" '
            "and fsale_date='" + today + "' and fzcounter='2251'"
        )

        update_frecno = (
            '\n\nUPDATE FRECNO:\n'
            "update pos_sale set fzcounter='' where fcompanyid='" + cid + "' "
            "and ftermid='" + terminal + "' and fsale_date='date nga e set' and fzcounter='0'"
        )

        sql_delete = (
            '\n\nSQL Query Delete:\n'
            "DELETE from pos_sale where frecno not in ('', '', '',);\n"
            "DELETE from pos_sale_payment where frecno not in ('', '', '',);\n"
            "DELETE from pos_sale_product where frecno not in ('', '', '',);\n"
            "DELETE from pos_reading;\nDELETE from pos_reading_summary;"
        )

        sql_update = (
            '\n\nSQL Query Update:\n'
            'UPDATE pos_sale_payment\n'
            "SET frecno = \nCASE frecno\n    WHEN '' THEN ''\nELSE frecno\nEND;"
        )

        sql_transmit = (
            '\n\nSQL Transmit:\n'
            'update pos_sale set ftransmit="0000";\n'
            'update pos_sale_payment set ftransmit="0000";\n'
            'update pos_sale_product set ftransmit="0000";'
        )

        for_exception = '\n\nFOR EXCEPTION:\n' + link + '/appserv/app/batch/sys_get_exception.php'
        for_bypass = '\n\nFOR BYPASS:\n' + link + '/appserv/app/batch/sys_bypass_exception.php/'
        for_rereading = (
            '\n\nFOR REREADING:\n' + link +
            '/appserv/app/batch/fix/recreate_reading.php?fcompanyid=' + cid +
            '&ftermid=' + terminal + '&fsale_date=&fzcounter=&fend_date=&fcreate_flag=1'
        )

        return lacking + full_missing + excel + select_or + update_frecno + sql_delete + sql_update + sql_transmit + for_exception + for_bypass + for_rereading

    def _validate(self):
        if not self.company_id.text().strip():
            AppDialog("Validation Error", "Company ID cannot be empty.", self).exec()
            return False
        if not self.terminal.text().strip():
            AppDialog("Validation Error", "Terminal cannot be empty.", self).exec()
            return False
        if not self.year.text().strip():
            AppDialog("Validation Error", "Missing Year cannot be empty.", self).exec()
            return False
        if not self.month.text().strip():
            AppDialog("Validation Error", "Missing Month cannot be empty.", self).exec()
            return False
        return True

    def generate_and_save(self):
        if not self._validate():
            return
        content = self._build_content()
        self.preview.setPlainText(content)

        today_folder = datetime.now().strftime("%m-%d-%Y")
        os.makedirs(today_folder, exist_ok=True)
        default_path = os.path.join(today_folder, "Script.txt")

        path, _ = QtWidgets.QFileDialog.getSaveFileName(
            self, "Save Script", default_path, "Text Files (*.txt);;All Files (*)"
        )
        if not path:
            return

        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            logger.info("Missing Sales script saved: %s", path)
            AppDialog("Success", f"Script saved to:\n{path}", self).exec()
        except Exception:
            logger.exception("Failed to save Missing Sales script")
            AppDialog("Error", "Failed to save script. Check logs.", self).exec()

    def copy_to_clipboard(self):
        if not self._validate():
            return
        content = self._build_content()
        self.preview.setPlainText(content)
        QtWidgets.QApplication.clipboard().setText(content)
        self.preview.append("\n--- Copied to clipboard ---")


# ================= MAIN WINDOW =================


class ReportWorker(QThread):
        progress = pyqtSignal(str)  # signal to update status
        finished = pyqtSignal(str)  # signal when done
        report_saved = pyqtSignal(str)

        def __init__(self, work_item_id, client_name, location_type1, location_type2, user_data, template_path):
            super().__init__()
            self.work_item_id = work_item_id
            self.client_name = client_name
            self.location_type1 = location_type1
            self.location_type2 = location_type2
            self.user_data = user_data
            self.template_path = template_path

        def run(self):
            logger.info("=== ReportWorker started for work item %s ===", self.work_item_id)
            pat = self.user_data["pat"]
            my_display_name = self.user_data["display_name"]
            my_unique_name = self.user_data["unique_name"]

            self.progress.emit(f"Fetching work item {self.work_item_id}...")
            work_item_url = f"{org_url}/_apis/wit/workitems/{self.work_item_id}?api-version=5.1"
            try:
                logger.debug("GET %s", work_item_url)
                resp = requests.get(work_item_url, auth=HTTPBasicAuth("", pat))
                resp.raise_for_status()
                work_item_data = resp.json()
                logger.info("Work item fetched: id=%s, title=%r", work_item_data.get("id"), work_item_data.get("fields", {}).get("System.Title", "")[:80])
            except Exception as e:
                logger.exception("Failed to fetch work item %s", self.work_item_id)
                self.finished.emit(f"❌ Error fetching work item: {e}")
                return

            ticket_number = work_item_data.get("id")
            ticket_subject = work_item_data.get("fields", {}).get("System.Title", "")

            self.progress.emit("Fetching comments...")
            comments_url = f"{org_url}/_apis/wit/workItems/{self.work_item_id}/comments?api-version=5.1-preview.3"
            try:
                logger.debug("GET %s", comments_url)
                resp = requests.get(comments_url, auth=HTTPBasicAuth("", pat))
                resp.raise_for_status()
                comments_data = resp.json()
                total_comments = len(comments_data.get("comments", []))
                logger.info("Fetched %d comment(s) for work item %s", total_comments, self.work_item_id)
            except Exception as e:
                logger.exception("Failed to fetch comments for work item %s", self.work_item_id)
                self.finished.emit(f"❌ Error fetching comments: {e}")
                return

            my_comments = [
                c for c in comments_data.get("comments", [])
                if c.get("createdBy", {}).get("displayName") == my_display_name
                or c.get("createdBy", {}).get("uniqueName") == my_unique_name
            ]
            logger.debug("Matched %d comment(s) for user displayName=%r, uniqueName=%r", len(my_comments), my_display_name, my_unique_name)
            if not my_comments:
                logger.warning("No comments found for user %r on work item %s", my_display_name, self.work_item_id)
                self.finished.emit("❌ No comments found for your account.")
                return

            latest_comment = sorted(my_comments, key=lambda x: x.get("createdDate", ""))[-1]
            self.progress.emit("Parsing comment...")

            iso_date = latest_comment.get("createdDate", "")
            logger.debug("Latest comment date (ISO): %s", iso_date)

            if iso_date:
                try:
                    # Handles timestamps WITH milliseconds
                    dt = datetime.strptime(iso_date, "%Y-%m-%dT%H:%M:%S.%fZ")
                except ValueError:
                    # Handles timestamps WITHOUT milliseconds
                    dt = datetime.strptime(iso_date, "%Y-%m-%dT%H:%M:%SZ")

                comment_date = dt.strftime("%m/%d/%Y")
            else:
                comment_date = ""

            logger.debug("Formatted comment date: %s", comment_date)

            raw_html = latest_comment.get("text", "")
            logger.debug("Raw HTML comment length: %d chars", len(raw_html))

            soup_for_text = BeautifulSoup(raw_html or "", "html.parser")
            clean_text = soup_for_text.get_text("\n", strip=True)

            # Known field labels (checked in order so longer names match first)
            known_labels = ["Root Cause", "Preventive Action", "Next Step", "Next Steps", "Status"]
            label_to_key = {
                "Root Cause": "Root Cause",
                "Preventive Action": "Preventive Action",
                "Next Step": "Next Step",
                "Next Steps": "Next Step",
                "Status": "Status",
            }

            def extract_all_fields(text):
                """Parse fields from comment text. Handles both
                'Label: value on same line' and 'Label:' on its own line
                with value on the next line."""
                lines = text.splitlines()
                fields = {}
                current_key = None
                current_value_lines = []

                def flush():
                    nonlocal current_key, current_value_lines
                    if current_key is not None:
                        fields[current_key] = "\n".join(current_value_lines).strip()
                    current_key = None
                    current_value_lines = []

                for line in lines:
                    stripped = line.strip()
                    matched = False
                    for label in known_labels:
                        after = stripped[len(label):]
                        if after.startswith(":"):
                            flush()
                            current_key = label_to_key[label]
                            remainder = after[1:].strip()
                            if remainder:
                                current_value_lines.append(remainder)
                            matched = True
                            break
                    if not matched and current_key is not None:
                        current_value_lines.append(stripped)

                flush()
                return fields

            fields = extract_all_fields(clean_text)
            root_cause = fields.get("Root Cause", "")
            preventive_action = fields.get("Preventive Action", "")
            next_step = fields.get("Next Step", "")
            status = fields.get("Status", "")

            logger.info("Extracted fields — Root Cause: %d chars, Preventive Action: %d chars, Next Step: %d chars, Status: %r",
                        len(root_cause), len(preventive_action), len(next_step), status)

            self.raw_html = raw_html
            self.root_cause = root_cause
            self.preventive_action = preventive_action
            self.next_step = next_step
            self.status = status

            today_folder = datetime.now().strftime("%m-%d-%Y")
            os.makedirs(today_folder, exist_ok=True)
            output_filename = f"Ticket No. {self.work_item_id} - Status Report for {self.client_name}.docx"
            output_path = os.path.join(today_folder, output_filename)
            self.last_output_path = output_path
            logger.info("Output path: %s", output_path)

            self.progress.emit("Generating DOCX...")
            try:
                logger.info("Loading template: %s", self.template_path)
                if not os.path.isfile(self.template_path):
                    logger.error("Template file NOT FOUND: %s (cwd=%s)", self.template_path, os.getcwd())
                    self.finished.emit(f"❌ Template not found: {self.template_path}")
                    return
                doc = Document(self.template_path)
                logger.info("Template loaded successfully (%d paragraph(s), %d table(s))", len(doc.paragraphs), len(doc.tables))
                placeholder_map = {
                    "{{TICKETNUM}}": str(ticket_number),
                    "{{DATE}}": comment_date,
                    "{{TICKETCONTENT}}": ticket_subject,
                    "{{RCA}}": root_cause,
                    "{{PREVAC}}": preventive_action,
                    "{{NEXT}}": next_step,
                    "{{CURSTAT}}": status,
                    "{{REPORT_CONTENT}}": raw_html,
                    "{{OS}}": self.location_type1,
                    "{{FS}}": self.location_type2,
                    "{{CLIENT}}": self.client_name,
                    "{{DISPLAYN}}": self.user_data["display_name"]  # new placeholder
                }

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                # Plain text placeholders
                                for ph in ["{{TICKETNUM}}", "{{TICKETCONTENT}}", "{{DATE}}", "{{OS}}", "{{FS}}", "{{CLIENT}}", "{{RCA}}", "{{PREVAC}}", "{{NEXT}}", "{{CURSTAT}}", "{{DISPLAYN}}"]:
                                    if ph in paragraph.text:
                                        logger.debug("Replacing %s in table paragraph", ph)
                                        paragraph.text = paragraph.text.replace(ph, placeholder_map[ph])

                                # HTML content placeholder
                                if "{{REPORT_CONTENT}}" in paragraph.text:
                                    logger.debug("Inserting HTML content into {{REPORT_CONTENT}} placeholder")
                                    for run in paragraph.runs:
                                        run.text = ""
                                    add_html_content_to_paragraph(paragraph, placeholder_map["{{REPORT_CONTENT}}"], pat)

                doc.save(output_path)
                logger.info("Report saved successfully: %s", output_path)
                self.report_saved.emit(output_path)
                self.finished.emit(f"✅ Report saved for ticket number: {self.work_item_id}")
            except Exception as e:
                logger.exception("Error generating report for work item %s", self.work_item_id)
                self.finished.emit(f"❌ Error generating report: {e}")

class PotTicketDataWorker(QThread):
    """
    Fetches the latest comment (by the logged-in user) for a given
    TFS work item, so it can be inserted into {{TICKETDATA}}.
    """
    progress = pyqtSignal(str)
    finished = pyqtSignal(str)   # emits raw_html on success, "" on failure
    error = pyqtSignal(str)      # emits error message on failure

    def __init__(self, work_item_id, user_data):
        super().__init__()
        self.work_item_id = work_item_id
        self.user_data = user_data
        self.raw_html = ""

    def run(self):
        logger.info("=== PotTicketDataWorker started for work item %s ===", self.work_item_id)
        import requests
        from requests.auth import HTTPBasicAuth

        pat = self.user_data["pat"]
        my_display_name = self.user_data["display_name"]
        my_unique_name = self.user_data.get("unique_name", "")

        self.progress.emit(f"Fetching comments for work item {self.work_item_id}...")
        comments_url = f"{org_url}/_apis/wit/workItems/{self.work_item_id}/comments?api-version=5.1-preview.3"

        try:
            logger.debug("GET %s", comments_url)
            resp = requests.get(comments_url, auth=HTTPBasicAuth("", pat))
            resp.raise_for_status()
            comments_data = resp.json()
            total = len(comments_data.get("comments", []))
            logger.info("Fetched %d comment(s) for POT work item %s", total, self.work_item_id)
        except Exception as e:
            logger.exception("Failed to fetch comments for POT work item %s", self.work_item_id)
            self.error.emit(f"❌ Error fetching comments: {e}")
            return

        my_comments = [
            c for c in comments_data.get("comments", [])
            if c.get("createdBy", {}).get("displayName") == my_display_name
            or c.get("createdBy", {}).get("uniqueName") == my_unique_name
        ]
        logger.debug("Matched %d comment(s) for POT user %r", len(my_comments), my_display_name)

        if not my_comments:
            logger.warning("No comments found for user %r on POT work item %s", my_display_name, self.work_item_id)
            self.error.emit("❌ No comments found for your account on this ticket.")
            return

        latest_comment = sorted(my_comments, key=lambda x: x.get("createdDate", ""))[-1]
        self.raw_html = latest_comment.get("text", "")
        logger.info("POT comment fetched: %d chars of HTML", len(self.raw_html))

        self.finished.emit(self.raw_html)

class LoginWindow(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Report — Marilag")
        self.setMinimumSize(460, 600)
        self.resize(460, 600)
        self.setObjectName("Root")
        self.apply_style()
        self.main_layout = QtWidgets.QVBoxLayout(self)
        self.main_layout.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)

        self.show_login()

    def _theme_colors(self):
        """Return a dict of theme-aware colors for inline widget styling."""
        if is_windows_dark_mode():
            return {
                "outline_border": "rgba(255,255,255,0.18)",
                "outline_text": "#e6eef5",
                "outline_hover_border": "#ed3b3b",
                "outline_hover_text": "#ffffff",
                "cancel_hover_border": "#98a0a6",
                "divider": "rgba(255,255,255,0.08)",
                "field_bg": "#15212c",
                "field_border": "#2a3a48",
                "field_text": "#e6eef5",
                "field_focus_border": "#ed3b3b",
                "section_label": "#ed3b3b",
                "field_label": "#98a0a6",
                "back_color": "#98a0a6",
                "back_hover": "#ff4d4d",
                "output_bg": "#111a22",
                "output_text": "#e6eef5",
            }
        else:
            return {
                "outline_border": "#d1d5db",
                "outline_text": "#1f2933",
                "outline_hover_border": "#ed3b3b",
                "outline_hover_text": "#ed3b3b",
                "cancel_hover_border": "#6b7280",
                "divider": "#e5e7eb",
                "field_bg": "#ffffff",
                "field_border": "#d1d5db",
                "field_text": "#111827",
                "field_focus_border": "#ed3b3b",
                "section_label": "#ed3b3b",
                "field_label": "#374151",
                "back_color": "#6b7280",
                "back_hover": "#ed3b3b",
                "output_bg": "#ffffff",
                "output_text": "#111827",
            }

    def set_page_size(self, width, height):
        """
        Resize the window for the current page. Called at the top of
        each show_xxx_page() method so different pages can use
        different window sizes (e.g. POT's wider two-column form).
        """
        self.setMinimumSize(width, height)
        self.resize(width, height)

    def apply_style(self):
        if is_windows_dark_mode():

            # DARK THEME
            self.setStyleSheet("""
            #Root {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1,
                    stop:0 #07101a, stop:1 #07121a);
                font-family: Inter, Segoe UI, Arial;
                color: #e6eef5;
            }

            QLabel#Title {
                font-size: 22px;
                font-weight: 700;
                color: #f8fbfc;
            }

            QLabel#Subtitle {
                font-size: 13px;
                color: #98a0a6;
            }

            QWidget#Card {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(255,255,255,0.05);
                border-radius: 14px;
            }

            QLineEdit {
                background: rgba(255,255,255,0.02);
                border: 1px solid rgba(255,255,255,0.06);
                border-radius: 10px;
                padding: 12px;
                font-size: 14px;
                color: #eaf6fb;
            }

            QLineEdit:focus {
                border: 1px solid #ed3b3b;
            }

            QPushButton {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1,
                    stop:0 #ed3b3b, stop:1 #c62e2e);
                border: none;
                border-radius: 10px;
                padding: 12px;
                font-size: 15px;
                font-weight: 700;
                color: white;
            }

            QPushButton:hover {
                opacity: 0.92;
            }

            QRadioButton {
                color: #e6eef5;
                spacing: 6px;
            }

            QRadioButton::indicator {
                width: 14px;
                height: 14px;
                border-radius: 7px;
                border: 2px solid #6b7280;
                background-color: transparent;
            }

            QRadioButton::indicator:checked {
                background-color: #DA3535;
                border: 2px solid #c62e2e;
            }

            QComboBox {
                background: rgba(255,255,255,0.02);
                border: 1px solid rgba(255,255,255,0.06);
                border-radius: 10px;
                padding: 12px 36px 12px 12px;
                font-size: 14px;
                color: #e6eef5;
                min-height: 20px;
            }

            QComboBox:focus {
                border: 1px solid #ed3b3b;
            }

            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 32px;
                border: none;
                border-left: 1px solid rgba(255,255,255,0.06);
                border-top-right-radius: 10px;
                border-bottom-right-radius: 10px;
            }

            QComboBox::down-arrow {
                width: 10px;
                height: 10px;
            }

            QComboBox QAbstractItemView {
                background: #15212c;
                border: 1px solid #2a3a48;
                border-radius: 8px;
                padding: 4px;
                color: #e6eef5;
                selection-background-color: #1e3a50;
                selection-color: #ffffff;
                outline: none;
            }

            QComboBox QAbstractItemView::item {
                padding: 8px 12px;
                border-radius: 4px;
                min-height: 24px;
            }

            QComboBox QAbstractItemView::item:hover {
                background: #1e3a50;
                color: #ffffff;
            }

            QDateEdit {
                background: rgba(255,255,255,0.02);
                border: 1px solid rgba(255,255,255,0.06);
                border-radius: 10px;
                padding: 12px 36px 12px 12px;
                font-size: 14px;
                color: #e6eef5;
                min-height: 20px;
            }

            QDateEdit:focus {
                border: 1px solid #ed3b3b;
            }

            QDateEdit::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 32px;
                border: none;
                border-left: 1px solid rgba(255,255,255,0.06);
                border-top-right-radius: 10px;
                border-bottom-right-radius: 10px;
            }

            QDateEdit::down-arrow {
                width: 0;
                height: 0;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 6px solid #98a0a6;
            }
            """)

        else:

            # LIGHT THEME
            self.setStyleSheet("""
            #Root {
                background: #f5f7fa;
                font-family: Inter, Segoe UI, Arial;
                color: #1f2933;
            }

            QLabel#Title {
                font-size: 22px;
                font-weight: 700;
                color: #111827;
            }

            QLabel#Subtitle {
                font-size: 13px;
                color: #6b7280;
            }

            QWidget#Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 14px;
            }

            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 10px;
                padding: 12px;
                font-size: 14px;
                color: #111827;
            }

            QLineEdit:focus {
                border: 1px solid #ed3b3b;
            }

            QPushButton {
                background: #ed3b3b;
                border: none;
                border-radius: 10px;
                padding: 12px;
                font-size: 15px;
                font-weight: 700;
                color: white;
            }

            QPushButton:hover {
                background: #c62e2e;
            }

            QRadioButton {
                color: #374151;
                spacing: 6px;
            }

            QRadioButton::indicator {
                width: 14px;
                height: 14px;
                border-radius: 7px;
                border: 2px solid #9ca3af;
                background-color: transparent;
            }

            QRadioButton::indicator:checked {
                background-color: #ed3b3b;
                border: 2px solid #c62e2e;
            }

            QComboBox {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 10px;
                padding: 12px 36px 12px 12px;
                font-size: 14px;
                color: #111827;
                min-height: 20px;
            }

            QComboBox:focus {
                border: 1px solid #ed3b3b;
            }

            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 32px;
                border: none;
                border-left: 1px solid #d1d5db;
                border-top-right-radius: 10px;
                border-bottom-right-radius: 10px;
            }

            QComboBox::down-arrow {
                width: 10px;
                height: 10px;
            }

            QComboBox QAbstractItemView {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 8px;
                padding: 4px;
                color: #111827;
                selection-background-color: #f0f0f0;
                selection-color: #111827;
                outline: none;
            }

            QComboBox QAbstractItemView::item {
                padding: 8px 12px;
                border-radius: 4px;
                min-height: 24px;
            }

            QComboBox QAbstractItemView::item:hover {
                background: #f3f4f6;
                color: #111827;
            }

            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 10px;
                padding: 12px 36px 12px 12px;
                font-size: 14px;
                color: #111827;
                min-height: 20px;
            }

            QDateEdit:focus {
                border: 1px solid #ed3b3b;
            }

            QDateEdit::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 32px;
                border: none;
                border-left: 1px solid #d1d5db;
                border-top-right-radius: 10px;
                border-bottom-right-radius: 10px;
            }

            QDateEdit::down-arrow {
                width: 0;
                height: 0;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 6px solid #6b7280;
            }
            """)

    #DARK THEME
    

    def confirm_generate_report(self):
        dialog = QtWidgets.QMessageBox(self)
        dialog.setWindowTitle("Confirm Action")
        dialog.setText("Do you want to create a status report?")
        dialog.setIcon(QtWidgets.QMessageBox.Icon.Question)

        dialog.setStandardButtons(
            QtWidgets.QMessageBox.StandardButton.Yes |
            QtWidgets.QMessageBox.StandardButton.No
        )

        if is_windows_dark_mode():
            dialog.setStyleSheet("""
            QMessageBox {
                background: #07101a;
                font-family: Inter, Segoe UI, Arial;
                color: #e6eef5;
            }
            QLabel {
                color: #e6eef5;
                background: transparent;
            }
            QPushButton {
                background: qlineargradient(x1:0,y1:0,x2:0,y2:1,
                    stop:0 #ed3b3b, stop:1 #c62e2e);
                color: white;
                border-radius: 8px;
                padding: 6px 16px;
                min-width: 60px;
            }
            QPushButton:hover {
                background: #c62e2e;
            }
            """)
        else:
            dialog.setStyleSheet("""
            QMessageBox {
                background: #f5f7fa;
                font-family: Inter, Segoe UI, Arial;
                color: #1f2933;
            }
            QLabel {
                color: #111827;
                background: transparent;
            }
            QPushButton {
                background: #ed3b3b;
                color: white;
                border-radius: 8px;
                padding: 6px 16px;
                min-width: 60px;
            }
            QPushButton:hover {
                background: #c62e2e;
            }
            """)

        return dialog.exec() == QtWidgets.QMessageBox.StandardButton.Yes



    def clear_layout(self):
        while self.main_layout.count():
            item = self.main_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()

    def build_card(self, title_text, subtitle_text, button_text, handler, pat=False, display=False, screenshot=None):
        card = QtWidgets.QWidget()
        card.setObjectName("Card")
        card.setFixedWidth(360)
        layout = QtWidgets.QVBoxLayout(card)
        layout.setContentsMargins(28, 28, 28, 28)
        layout.setSpacing(16)

        # Optional screenshot at top
        if screenshot:
            pixmap = QtGui.QPixmap(screenshot)
            if not pixmap.isNull():
                pixmap = pixmap.scaledToWidth(320, QtCore.Qt.TransformationMode.SmoothTransformation)
                img_label = QtWidgets.QLabel()
                img_label.setPixmap(pixmap)
                img_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
                img_label.setContentsMargins(0, 0, 0, 12)
                layout.addWidget(img_label)
            else:
                print(f"Warning: Screenshot not found at {screenshot}")

        # Title
        title = QtWidgets.QLabel(title_text)
        title.setObjectName("Title")
        layout.addWidget(title)

        # Subtitle / description with word wrap and selectable text
        subtitle = QtWidgets.QLabel(subtitle_text)
        subtitle.setObjectName("Subtitle")
        subtitle.setWordWrap(True)  # Wrap long text
        subtitle.setTextInteractionFlags(QtCore.Qt.TextInteractionFlag.TextSelectableByMouse)  # Make text copyable
        layout.addWidget(subtitle)

        layout.addSpacing(6)

        # Input fields
        if pat:
            self.pat_input = QtWidgets.QLineEdit()
            self.pat_input.setPlaceholderText("Enter your PAT token")
            layout.addWidget(self.pat_input)
        elif display:
            self.display_input = QtWidgets.QLineEdit()
            self.display_input.setPlaceholderText("Enter Azure Display Name")
            layout.addWidget(self.display_input)
        else:
            self.username = QtWidgets.QLineEdit()
            self.username.setPlaceholderText("Email")
            self.password = QtWidgets.QLineEdit()
            self.password.setPlaceholderText("Password")
            self.password.setEchoMode(QtWidgets.QLineEdit.EchoMode.Password)
            layout.addWidget(self.username)
            layout.addWidget(self.password)

        layout.addSpacing(10)

        button = QtWidgets.QPushButton(button_text)
        button.clicked.connect(handler)
        button.setFixedHeight(40)

        # Store current page button so we can change text like "Creating Account..."
        self.action_button = button

        layout.addWidget(button)

        for field in card.findChildren(QtWidgets.QLineEdit):
            field.returnPressed.connect(button.click)

        self.main_layout.addWidget(card, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)



    # ----------------- PAGES -----------------
    def show_create_account(self):
        self.clear_layout()
        self.build_card(
            "Create your account",
            "Set a username and password for this app",
            "Next",
            self.next_to_pat
        )

    def next_to_pat(self):
        u = self.username.text().strip()
        p = self.password.text()
        if len(u) < 3 or len(p) < 6:
            AppDialog("Error", "Username must be 3+ chars\nPassword must be 6+ chars", self).exec()
            return
        self.temp_username = u
        self.temp_password = p
        self.show_pat_page()

    def show_pat_page(self):
        self.clear_layout()
        self.build_card(
            "Create PAT Token",
            "Create PAT by accessing the link below: \nhttps://tfs.alliancewebpos.com/tfs/WebPOSCollection/_usersSettings/tokens",
            "Next",           # button text
            self.save_pat,    # handler
            pat=True,
            screenshot="asset/picture01.png"
        )

    def save_pat(self):
        pat = self.pat_input.text().strip()
        if len(pat) < 6:
            AppDialog("Error", "PAT must be at least 6 characters", self).exec()
            return
        self.temp_pat = pat
        self.show_display_name_page()

    def show_display_name_page(self):
        self.clear_layout()
        self.build_card(
            "Azure Display Name",
            "Enter the display name used in Azure for your account",
            "Finish Setup",
            self.save_display_name,
            display=True,
            screenshot="asset/picture1.png"  # <-- Put your screenshot path here
        )

    # Replace the save_display_name method with this:
    def save_display_name(self):
        display_name = self.display_input.text().strip()
        if len(display_name) < 1:
            AppDialog("Error", "Display Name cannot be empty", self).exec()
            return

        self.temp_display_name = display_name  # store for next step
        self.show_unique_name_page()  # move to unique name page instead of saving

    def show_unique_name_page(self):
        self.clear_layout()
        self.build_card(
            "Unique Name",
            "Enter a unique name for this account (used for identification in the app)",
            "Finish Setup",
            self.save_unique_name,
            display=True,  # reuse the input field type
            screenshot="asset/picture2.png"  # optional screenshot for professionalism
        )

    def save_unique_name(self):
        unique_name = self.display_input.text().strip()

        if len(unique_name) < 1:
            AppDialog("Error", "Unique Name cannot be empty", self).exec()
            return

        logger.info("Account creation attempt for email=%r, display_name=%r", self.temp_username, self.temp_display_name)
        try:
            self.action_button.setEnabled(False)
            self.action_button.setText("Creating Account...")
            QApplication.processEvents()

            auth, db = _get_firebase()
            user = auth.create_user_with_email_and_password(
                self.temp_username,
                self.temp_password
            )

            uid = user["localId"]
            logger.info("Firebase user created, uid=%s", uid)

            self.action_button.setText("Saving User Data...")
            QApplication.processEvents()

            db.collection("users").document(uid).set({
                "display_name": self.temp_display_name,
                "unique_name": unique_name,
                "pat": self.temp_pat
            })
            logger.info("Firestore profile saved for uid=%s", uid)

            # Try to fetch PAT expiry and save it
            self.action_button.setText("Checking PAT expiry...")
            QApplication.processEvents()
            pat_expiry = _fetch_pat_expiry(self.temp_pat)
            if pat_expiry:
                db.collection("users").document(uid).update({
                    "pat_expiry": pat_expiry
                })
                logger.info("PAT expiry saved to Firestore: %s", pat_expiry)
            else:
                logger.info("PAT expiry could not be determined (will retry on login)")

            self.action_button.setText("Success!")
            QApplication.processEvents()

            AppDialog("Success", "Account created successfully!", self).exec()
            self.show_login()

        except Exception as e:
            self.action_button.setEnabled(True)
            self.action_button.setText("Finish Setup")

            error_message = str(e)
            logger.warning("Account creation failed for %r: %s", self.temp_username, error_message)

            if "EMAIL_EXISTS" in error_message:
                friendly_message = "This email is already registered.\nPlease login instead."
            elif "INVALID_EMAIL" in error_message:
                friendly_message = "Please enter a valid email address."
            elif "WEAK_PASSWORD" in error_message:
                friendly_message = "Password is too weak.\nPlease use at least 6 characters."
            else:
                friendly_message = "Registration failed.\nPlease try again later."

            AppDialog("Registration Error", friendly_message, self).exec()

    def show_main_menu_page(self):
        self.clear_layout()
        self.set_page_size(460, 620)

        card = QtWidgets.QWidget()
        card.setObjectName("Card")
        card.setFixedWidth(380)

        layout = QtWidgets.QVBoxLayout(card)
        layout.setContentsMargins(36, 36, 36, 32)
        layout.setSpacing(0)

        tc = self._theme_colors()

        # --- Top bar: Eyebrow + Settings gear icon ---
        top_bar = QtWidgets.QHBoxLayout()
        top_bar.setSpacing(0)

        eyebrow = QtWidgets.QLabel("WORKSPACE")
        eyebrow.setStyleSheet("""
            color: #ed3b3b;
            font-size: 11px;
            font-weight: 700;
            letter-spacing: 2px;
        """)
        top_bar.addWidget(eyebrow)
        top_bar.addStretch()

        settings_icon = QtWidgets.QToolButton()
        settings_icon.setText("\u2699")
        settings_icon.setToolTip("Settings")
        settings_icon.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        settings_icon.clicked.connect(self.show_settings_page)
        settings_icon.setStyleSheet(f"""
            QToolButton {{
                border: none;
                font-size: 18px;
                color: {tc['back_color']};
                padding: 0;
            }}
            QToolButton:hover {{
                color: {tc['back_hover']};
            }}
        """)
        top_bar.addWidget(settings_icon)

        layout.addLayout(top_bar)
        layout.addSpacing(8)

        # --- Title ---
        title = QtWidgets.QLabel("What are you creating?")
        title.setObjectName("Title")
        title.setWordWrap(True)
        layout.addWidget(title)
        layout.addSpacing(6)

        # --- Subtitle ---
        subtitle = QtWidgets.QLabel("Pick a document type to get started.")
        subtitle.setObjectName("Subtitle")
        subtitle.setWordWrap(True)
        layout.addWidget(subtitle)

        layout.addSpacing(28)

        # --- Divider ---
        divider = QtWidgets.QFrame()
        divider.setFixedHeight(1)
        divider.setStyleSheet(f"background: {tc['divider']}; border: none;")
        layout.addWidget(divider)
        layout.addSpacing(24)

        # --- Primary action: Create Status Report ---
        status_btn = QtWidgets.QPushButton("Create Status Report")
        status_btn.setFixedHeight(46)
        status_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        status_btn.clicked.connect(self.show_report_page)
        layout.addWidget(status_btn)

        layout.addSpacing(12)

        # --- Secondary action: Create POT (outlined style) ---
        pot_btn = QtWidgets.QPushButton("Create POT")
        pot_btn.setFixedHeight(46)
        pot_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        pot_btn.clicked.connect(self.show_pot_page)
        pot_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                border: 1.5px solid {tc['outline_border']};
                border-radius: 10px;
                padding: 12px;
                font-size: 15px;
                font-weight: 600;
                color: {tc['outline_text']};
            }}
            QPushButton:hover {{
                border: 1.5px solid {tc['outline_hover_border']};
                color: {tc['outline_hover_text']};
            }}
        """)
        layout.addWidget(pot_btn)

        layout.addSpacing(12)

        # --- Tertiary action: Tools (outlined style) ---
        tools_btn = QtWidgets.QPushButton("Tools")
        tools_btn.setFixedHeight(46)
        tools_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        tools_btn.clicked.connect(self.show_tools_page)
        tools_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                border: 1.5px solid {tc['outline_border']};
                border-radius: 10px;
                padding: 12px;
                font-size: 15px;
                font-weight: 600;
                color: {tc['outline_text']};
            }}
            QPushButton:hover {{
                border: 1.5px solid {tc['outline_hover_border']};
                color: {tc['outline_hover_text']};
            }}
        """)
        layout.addWidget(tools_btn)

        # --- Footer: signed in as ---
        if hasattr(self, "current_user") and self.current_user.get("display_name"):
            layout.addSpacing(28)
            footer_divider = QtWidgets.QFrame()
            footer_divider.setFixedHeight(1)
            footer_divider.setStyleSheet(f"background: {tc['divider']}; border: none;")
            layout.addWidget(footer_divider)
            layout.addSpacing(14)

            footer = QtWidgets.QLabel(f"Signed in as {self.current_user['display_name']}")
            footer.setStyleSheet("color: #6b7280; font-size: 12px;")
            footer.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(footer)

            build_label = QtWidgets.QLabel(f"Release v{APP_VERSION}")
            build_label.setStyleSheet("color: #6b7280; font-size: 11px;")
            build_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(build_label)

        self.main_layout.addWidget(
            card,
            alignment=QtCore.Qt.AlignmentFlag.AlignCenter
        )


    def show_settings_page(self):
        self.clear_layout()
        self.set_page_size(460, 520)

        card = QtWidgets.QWidget()
        card.setObjectName("Card")
        card.setFixedWidth(380)

        layout = QtWidgets.QVBoxLayout(card)
        layout.setContentsMargins(36, 36, 36, 32)
        layout.setSpacing(0)

        tc = self._theme_colors()

        # --- Top Bar (Back Arrow) ---
        top_bar = QtWidgets.QHBoxLayout()
        back_btn = QtWidgets.QToolButton()
        back_btn.setText("\u2190")
        back_btn.setToolTip("Back to Main Menu")
        back_btn.clicked.connect(self.show_main_menu_page)
        back_btn.setStyleSheet(f"""
            QToolButton {{
                border: none;
                font-size: 18px;
                font-weight: bold;
                color: {tc['back_color']};
            }}
            QToolButton:hover {{
                color: {tc['back_hover']};
            }}
        """)
        back_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        top_bar.addWidget(back_btn)
        top_bar.addStretch()
        layout.addLayout(top_bar)
        layout.addSpacing(4)

        # --- Title ---
        title = QtWidgets.QLabel("Settings")
        title.setObjectName("Title")
        layout.addWidget(title)
        layout.addSpacing(6)

        subtitle = QtWidgets.QLabel("Manage your account and PAT settings.")
        subtitle.setObjectName("Subtitle")
        subtitle.setWordWrap(True)
        layout.addWidget(subtitle)

        layout.addSpacing(28)

        # --- Divider ---
        divider = QtWidgets.QFrame()
        divider.setFixedHeight(1)
        divider.setStyleSheet(f"background: {tc['divider']}; border: none;")
        layout.addWidget(divider)
        layout.addSpacing(20)

        # --- PAT Expiry Section ---
        section_label = QtWidgets.QLabel("PAT EXPIRY")
        section_label.setStyleSheet(f"""
            color: {tc['section_label']};
            font-size: 11px;
            font-weight: 700;
            letter-spacing: 1.5px;
        """)
        layout.addWidget(section_label)
        layout.addSpacing(12)

        field_label = QtWidgets.QLabel("PAT Token Expiry Date")
        field_label.setStyleSheet(f"color: {tc['field_label']}; font-size: 12px; font-weight: 600;")
        layout.addWidget(field_label)
        layout.addSpacing(6)

        self.settings_expiry_input = QtWidgets.QDateEdit()
        self.settings_expiry_input.setCalendarPopup(True)
        self.settings_expiry_input.setDate(QtCore.QDate.currentDate().addYears(1))
        self.settings_expiry_input.setStyleSheet(f"""
            QDateEdit {{
                background-color: {tc['field_bg']};
                border: 1px solid {tc['field_border']};
                border-radius: 8px;
                padding: 9px 36px 9px 10px;
                color: {tc['field_text']};
                font-size: 13px;
            }}
            QDateEdit:focus {{
                border: 1px solid {tc['field_focus_border']};
            }}
            QDateEdit::drop-down {{
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 28px;
                border: none;
                border-left: 1px solid {tc['field_border']};
                border-top-right-radius: 8px;
                border-bottom-right-radius: 8px;
            }}
            QDateEdit::down-arrow {{
                width: 0;
                height: 0;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 6px solid {tc['field_label']};
            }}
        """)
        layout.addWidget(self.settings_expiry_input)

        layout.addSpacing(8)

        # Show current saved expiry
        current_expiry = self.current_user.get("pat_expiry", "") if hasattr(self, "current_user") else ""
        if current_expiry:
            expiry_display = QtWidgets.QLabel(f"Current: {current_expiry}")
            expiry_display.setStyleSheet("color: #6b7280; font-size: 11px;")
            layout.addWidget(expiry_display)
            # Try to pre-fill the date picker with the saved value
            try:
                dt = datetime.fromisoformat(current_expiry.replace("Z", "+00:00"))
                self.settings_expiry_input.setDate(QtCore.QDate(dt.year, dt.month, dt.day))
            except Exception:
                pass
        else:
            expiry_display = QtWidgets.QLabel("No expiry date saved yet")
            expiry_display.setStyleSheet("color: #6b7280; font-size: 11px;")
            layout.addWidget(expiry_display)

        layout.addSpacing(24)

        # --- Save Button ---
        btn_layout = QtWidgets.QHBoxLayout()
        save_btn = QtWidgets.QPushButton("Save")
        save_btn.setFixedHeight(42)
        save_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        save_btn.clicked.connect(self.save_pat_expiry)
        save_btn.setFixedWidth(120)

        btn_layout.addStretch()
        btn_layout.addWidget(save_btn)
        layout.addLayout(btn_layout)

        layout.addSpacing(12)

        # --- Status label ---
        self.settings_status = QtWidgets.QLabel("")
        self.settings_status.setWordWrap(True)
        if is_windows_dark_mode():
            self.settings_status.setStyleSheet("color: #4ade80; font-weight: 600; font-size: 12.5px;")
        else:
            self.settings_status.setStyleSheet("color: #16a34a; font-weight: 600; font-size: 12.5px;")
        layout.addWidget(self.settings_status)

        self.main_layout.addWidget(card, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)


    def save_pat_expiry(self):
        expiry_date = self.settings_expiry_input.date()
        expiry_str = expiry_date.toString("yyyy-MM-dd") + "T00:00:00Z"

        uid = self.current_user.get("uid", "") if hasattr(self, "current_user") else ""
        if not uid:
            self.settings_status.setText("Not logged in.")
            return

        logger.info("Saving PAT expiry: %s", expiry_str)
        try:
            auth, db = _get_firebase()

            db.collection("users").document(uid).update({
                "pat_expiry": expiry_str
            })
            self.current_user["pat_expiry"] = expiry_str
            logger.info("PAT expiry saved to Firestore for uid=%s", uid)
            self.settings_status.setText(f"Saved! Expiry: {expiry_date.toString('MM/dd/yyyy')}")
        except Exception:
            logger.exception("Failed to save PAT expiry")
            self.settings_status.setText("Failed to save. Check logs.")


    # ================= TOOLS PAGE =================
    def show_tools_page(self):
        self.clear_layout()
        self.set_page_size(480, 620)

        tc = self._theme_colors()

        outer = QtWidgets.QVBoxLayout()
        outer.setContentsMargins(24, 20, 24, 20)
        outer.setSpacing(0)

        # --- Top Bar (Back Arrow) ---
        top_bar = QtWidgets.QHBoxLayout()
        back_btn = QtWidgets.QToolButton()
        back_btn.setText("←")
        back_btn.setToolTip("Back to Main Menu")
        back_btn.clicked.connect(self.show_main_menu_page)
        back_btn.setStyleSheet(f"""
            QToolButton {{
                border: none;
                font-size: 18px;
                font-weight: bold;
                color: {tc['back_color']};
            }}
            QToolButton:hover {{
                color: {tc['back_hover']};
            }}
        """)
        back_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        top_bar.addWidget(back_btn)
        top_bar.addStretch()
        outer.addLayout(top_bar)
        outer.addSpacing(4)

        # --- Title ---
        title = QtWidgets.QLabel("Tools")
        title.setObjectName("Title")
        outer.addWidget(title)
        outer.addSpacing(2)

        subtitle = QtWidgets.QLabel(f"{len(TOOL_REGISTRY)} tools available")
        subtitle.setObjectName("Subtitle")
        outer.addWidget(subtitle)
        outer.addSpacing(12)

        # --- Search bar ---
        self.tools_search = QtWidgets.QLineEdit()
        self.tools_search.setPlaceholderText("Search tools...")
        self.tools_search.setClearButtonEnabled(True)
        self.tools_search.setStyleSheet(f"""
            QLineEdit {{
                background-color: {tc['field_bg']};
                border: 1px solid {tc['field_border']};
                border-radius: 8px;
                padding: 8px 10px;
                color: {tc['field_text']};
                font-size: 13px;
            }}
            QLineEdit:focus {{
                border: 1px solid {tc['field_focus_border']};
            }}
        """)
        outer.addWidget(self.tools_search)
        outer.addSpacing(12)

        # --- Scrollable list ---
        scroll_area = QtWidgets.QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll_area.setStyleSheet("""
            QScrollArea { border: none; background: transparent; }
            QScrollBar:vertical { width: 5px; background: transparent; }
            QScrollBar::handle:vertical { background: rgba(128,128,128,0.3); border-radius: 2px; min-height: 20px; }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
        """)

        scroll_widget = QtWidgets.QWidget()
        scroll_widget.setStyleSheet("background: transparent;")
        self.tools_list_layout = QtWidgets.QVBoxLayout(scroll_widget)
        self.tools_list_layout.setSpacing(6)
        self.tools_list_layout.setContentsMargins(2, 2, 2, 2)

        self._populate_tools_list(tc)

        scroll_area.setWidget(scroll_widget)
        outer.addWidget(scroll_area)

        container = QtWidgets.QWidget()
        container.setLayout(outer)
        self.main_layout.addWidget(container)

        # --- Connect search to filter ---
        self.tools_search.textChanged.connect(lambda text: self._filter_tools(text, tc))

    def _populate_tools_list(self, tc, filter_text=""):
        """Populate the tools list from TOOL_REGISTRY. Auto-fetches all tools."""
        while self.tools_list_layout.count():
            item = self.tools_list_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()

        found = 0
        for tool in TOOL_REGISTRY:
            if filter_text.lower() not in tool["name"].lower() and \
               filter_text.lower() not in tool.get("description", "").lower() and \
               filter_text.lower() not in tool.get("category", "").lower():
                continue

            row = self._create_tool_row(tool, tc)
            self.tools_list_layout.addWidget(row)
            found += 1

        if found == 0:
            no_results = QtWidgets.QLabel("No tools found")
            no_results.setStyleSheet(f"color: {tc['back_color']}; font-size: 13px; padding: 20px;")
            no_results.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.tools_list_layout.addWidget(no_results)

        self.tools_list_layout.addStretch()

    def _filter_tools(self, text, tc):
        """Re-populate list when search text changes."""
        self._populate_tools_list(tc, filter_text=text)

    def _create_tool_row(self, tool, tc):
        """Create a compact horizontal row for a single tool."""
        row = QtWidgets.QFrame()
        row.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        row.setMinimumHeight(56)
        row.setStyleSheet(f"""
            QFrame {{
                background-color: {tc['field_bg']};
                border: 1px solid {tc['field_border']};
                border-radius: 8px;
            }}
            QFrame:hover {{
                border: 1px solid {tc['field_focus_border']};
            }}
        """)

        row_layout = QtWidgets.QHBoxLayout(row)
        row_layout.setContentsMargins(14, 8, 14, 8)
        row_layout.setSpacing(12)

        # Icon
        icon_label = QtWidgets.QLabel(tool.get("icon", "🔧"))
        icon_label.setFixedWidth(28)
        icon_label.setStyleSheet("font-size: 18px; background: transparent; border: none;")
        row_layout.addWidget(icon_label)

        # Text block (name + description stacked)
        text_block = QtWidgets.QVBoxLayout()
        text_block.setSpacing(2)

        name_label = QtWidgets.QLabel(tool["name"])
        name_label.setStyleSheet(f"""
            font-size: 13px; font-weight: 600; color: {tc['field_text']};
            background: transparent; border: none;
        """)
        text_block.addWidget(name_label)

        desc_label = QtWidgets.QLabel(tool.get("description", ""))
        desc_label.setStyleSheet(f"""
            font-size: 11px; color: {tc['back_color']};
            background: transparent; border: none;
        """)
        desc_label.setWordWrap(True)
        text_block.addWidget(desc_label)

        row_layout.addLayout(text_block, stretch=1)

        # Arrow indicator
        arrow = QtWidgets.QLabel("›")
        arrow.setStyleSheet(f"""
            font-size: 18px; color: {tc['back_color']};
            background: transparent; border: none;
        """)
        row_layout.addWidget(arrow)

        # Click handler
        handler_name = tool.get("handler", "")
        row.mousePressEvent = lambda event, h=handler_name: self._on_tool_clicked(h)

        return row

    def _on_tool_clicked(self, handler_name):
        """Dispatch to the tool's handler method on this window."""
        if not handler_name:
            AppDialog("Coming Soon", "This tool is not yet implemented.", self).exec()
            return
        handler = getattr(self, handler_name, None)
        if handler and callable(handler):
            logger.info("Tool clicked: %s", handler_name)
            handler()
        else:
            AppDialog("Error", f"Handler '{handler_name}' not found.", self).exec()

    # ================= TOOL HANDLERS =================
    # Implement each tool's logic below.
    # The handler name must match the "handler" key in TOOL_REGISTRY.

    def run_tool_1(self):
        dialog = MissingSalesDialog(self)
        dialog.exec()

    def run_tool_2(self):
        """Tool 2 — TODO: implement"""
        AppDialog("Tool 2", "Tool 2 is not yet implemented.", self).exec()

    def run_tool_3(self):
        """Tool 3 — TODO: implement"""
        AppDialog("Tool 3", "Tool 3 is not yet implemented.", self).exec()

    def run_tool_4(self):
        """Tool 4 — TODO: implement"""
        AppDialog("Tool 4", "Tool 4 is not yet implemented.", self).exec()

    def run_tool_5(self):
        """Tool 5 — TODO: implement"""
        AppDialog("Tool 5", "Tool 5 is not yet implemented.", self).exec()

    def run_tool_6(self):
        """Tool 6 — TODO: implement"""
        AppDialog("Tool 6", "Tool 6 is not yet implemented.", self).exec()

    def run_tool_7(self):
        """Tool 7 — TODO: implement"""
        AppDialog("Tool 7", "Tool 7 is not yet implemented.", self).exec()

    def run_tool_8(self):
        """Tool 8 — TODO: implement"""
        AppDialog("Tool 8", "Tool 8 is not yet implemented.", self).exec()

    def run_tool_9(self):
        """Tool 9 — TODO: implement"""
        AppDialog("Tool 9", "Tool 9 is not yet implemented.", self).exec()

    def run_tool_10(self):
        """Tool 10 — TODO: implement"""
        AppDialog("Tool 10", "Tool 10 is not yet implemented.", self).exec()


    def show_login(self):
        self.clear_layout()
        self.set_page_size(460, 600)

        card = QtWidgets.QWidget()
        card.setObjectName("Card")
        card.setFixedWidth(360)

        layout = QtWidgets.QVBoxLayout(card)
        layout.setContentsMargins(28, 28, 28, 28)
        layout.setSpacing(16)

        title = QtWidgets.QLabel("Sign in")
        title.setObjectName("Title")
        layout.addWidget(title)

        subtitle = QtWidgets.QLabel("Enter your email and password to continue")
        subtitle.setObjectName("Subtitle")
        subtitle.setWordWrap(True)
        layout.addWidget(subtitle)

        self.username = QtWidgets.QLineEdit()
        self.username.setPlaceholderText("Email")

        self.password = QtWidgets.QLineEdit()
        self.password.setPlaceholderText("Password")
        self.password.setEchoMode(QtWidgets.QLineEdit.EchoMode.Password)

        layout.addWidget(self.username)
        layout.addWidget(self.password)

        self.login_btn = QtWidgets.QPushButton("Login")
        self.login_btn.clicked.connect(self.login)
        self.login_btn.setFixedHeight(40)
        layout.addWidget(self.login_btn)

        register_label = QtWidgets.QLabel(
            'Don\'t have an account? '
            '<a href="#">Create one</a>'
        )

        register_label.setObjectName("Subtitle")
        register_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)

        register_label.setTextFormat(QtCore.Qt.TextFormat.RichText)
        register_label.setTextInteractionFlags(
            QtCore.Qt.TextInteractionFlag.TextBrowserInteraction
        )

        register_label.setOpenExternalLinks(False)

        register_label.linkActivated.connect(
            lambda _: self.show_create_account()
        )

        layout.addWidget(register_label)

        build_label = QtWidgets.QLabel(f"Release v{APP_VERSION}")
        build_label.setStyleSheet("color: #6b7280; font-size: 11px;")
        build_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(build_label)

        self.username.returnPressed.connect(self.login_btn.click)
        self.password.returnPressed.connect(self.login_btn.click)

        self.main_layout.addWidget(
            card,
            alignment=QtCore.Qt.AlignmentFlag.AlignCenter
        )

    def login(self):
        email = self.username.text().strip()
        password = self.password.text()

        if not email or not password:
            AppDialog("Error", "Email and password cannot be empty.", self).exec()
            return

        logger.info("Login attempt for email=%r", email)
        try:
            self.login_btn.setEnabled(False)
            self.login_btn.setText("Signing in...")
            QApplication.processEvents()

            # Firebase Authentication login
            auth, db = _get_firebase()
            user = auth.sign_in_with_email_and_password(email, password)
            uid = user["localId"]
            logger.info("Firebase auth success, uid=%s", uid)

            self.login_btn.setText("Fetching Profile...")
            QApplication.processEvents()

            # Fetch user profile from Firestore
            doc = db.collection("users").document(uid).get()

            if not doc.exists:
                raise Exception("USER_PROFILE_NOT_FOUND")

            user_data = doc.to_dict()
            logger.info("Firestore profile loaded: display_name=%r", user_data.get("display_name"))

            # Store current user in memory
            self.current_user = {
                "email": email,
                "uid": uid,
                "pat": user_data.get("pat", ""),
                "display_name": user_data.get("display_name", ""),
                "unique_name": user_data.get("unique_name", ""),
                "pat_expiry": user_data.get("pat_expiry", ""),
            }

            # If pat_expiry is missing from Firestore, try to fetch and save it
            if not self.current_user["pat_expiry"] and self.current_user["pat"]:
                logger.info("pat_expiry missing in Firestore — fetching from TFS")
                self.login_btn.setText("Checking PAT expiry...")
                QApplication.processEvents()
                pat_expiry = _fetch_pat_expiry(self.current_user["pat"])
                if pat_expiry:
                    db.collection("users").document(uid).update({
                        "pat_expiry": pat_expiry
                    })
                    self.current_user["pat_expiry"] = pat_expiry
                    logger.info("PAT expiry saved to Firestore: %s", pat_expiry)
                else:
                    logger.info("PAT expiry could not be determined")

            AppDialog(
                "Welcome",
                f"Hello {self.current_user['display_name']}!\nYou have successfully logged in.",
                self
            ).exec()

            # self.show_report_page()
            self.show_main_menu_page()
        except Exception as e:
            self.login_btn.setEnabled(True)
            self.login_btn.setText("Login")

            error_message = str(e)
            logger.warning("Login failed for %r: %s", email, error_message)

            if "INVALID_LOGIN_CREDENTIALS" in error_message:
                friendly_message = "Invalid email or password."
            elif "EMAIL_NOT_FOUND" in error_message:
                friendly_message = "No account found with this email."
            elif "INVALID_PASSWORD" in error_message:
                friendly_message = "Incorrect password."
            elif "USER_DISABLED" in error_message:
                friendly_message = "This account has been disabled."
            elif "USER_PROFILE_NOT_FOUND" in error_message:
                friendly_message = "Account found, but profile data is missing."
            else:
                friendly_message = "Login failed.\nPlease try again later."

            AppDialog("Login Error", friendly_message, self).exec()

    def show_azure_report_page(self):
        # Clear the current layout
        self.clear_layout()

        card = QtWidgets.QWidget()
        card.setObjectName("Card")
        card.setFixedWidth(400)
        layout = QtWidgets.QVBoxLayout(card)
        layout.setContentsMargins(28, 28, 28, 28)
        layout.setSpacing(12)

        # Title
        title = QtWidgets.QLabel("Create Azure Status Report")
        title.setObjectName("Title")
        layout.addWidget(title)

        # Subtitle
        subtitle = QtWidgets.QLabel("Fill in Azure-specific details to continue")
        subtitle.setObjectName("Subtitle")
        subtitle.setWordWrap(True)
        layout.addWidget(subtitle)

        # Example input field
        self.azure_comment_input = QtWidgets.QLineEdit()
        self.azure_comment_input.setPlaceholderText("Additional comment for Azure report")
        layout.addWidget(self.azure_comment_input)

        # Continue button
        continue_btn = QtWidgets.QPushButton("Submit Azure Report")
        continue_btn.clicked.connect(self.submit_azure_report)
        layout.addWidget(continue_btn)

        self.main_layout.addWidget(card, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)


    def submit_azure_report(self):
        comment = self.azure_comment_input.text().strip()
        if not comment:
            AppDialog("Error", "Please enter a comment", self).exec()
            return

        # Here you can implement Azure report logic
        AppDialog("Success", "Azure report submitted!", self).exec()

        # Optionally return to report page
        # self.show_report_page()
        self.show_main_menu_page()
    

    def handle_report_finished(self, msg):
        self.output_area.append(msg)

        # Re-enable inputs
        self.generate_report_btn.setEnabled(True)
        self.generate_report_btn.setText("Generate Report")
        self.work_item_input.setEnabled(True)
        self.client_input.setEnabled(True)
        self.onsite_radio.setEnabled(True)
        self.offsite_radio.setEnabled(True)

        if msg.startswith("✅ Report saved"):
            # Show and enable the "Create Azure Status Report" button
            if hasattr(self, "create_azure_btn") and self.create_azure_btn is not None:
                self.create_azure_btn.setVisible(True)
                self.create_azure_btn.setEnabled(True)

            # Save report data for later use in PBI creation
            self.last_report_data = {
                "work_item_id": self.worker.work_item_id,
                "client_name": self.worker.client_name,
                "root_cause": getattr(self.worker, "root_cause", ""),
                "preventive_action": getattr(self.worker, "preventive_action", ""),
                "next_step": getattr(self.worker, "next_step", ""),
                "raw_html": getattr(self.worker, "raw_html", "")
            }

    def show_pot_page(self):
        self.clear_layout()
        self.set_page_size(700, 760)

        # Card container
        card = QtWidgets.QWidget()
        card.setObjectName("Card")
        card.setFixedWidth(620)

        layout = QtWidgets.QVBoxLayout(card)
        layout.setContentsMargins(32, 28, 32, 28)
        layout.setSpacing(0)

        tc = self._theme_colors()

        # ===== Top Bar (Back Arrow + Folder Icon) =====
        top_bar = QtWidgets.QHBoxLayout()
        back_btn = QtWidgets.QToolButton()
        back_btn.setText("←")
        back_btn.setToolTip("Back to Main Menu")
        back_btn.clicked.connect(self.show_main_menu_page)
        back_btn.setStyleSheet(f"""
            QToolButton {{
                border: none;
                font-size: 18px;
                font-weight: bold;
                color: {tc['back_color']};
            }}
            QToolButton:hover {{
                color: {tc['back_hover']};
            }}
        """)
        back_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)

        self.pot_folder_btn = QtWidgets.QToolButton()
        self.pot_folder_btn.setIcon(self.style().standardIcon(QtWidgets.QStyle.StandardPixmap.SP_DirOpenIcon))
        self.pot_folder_btn.setToolTip("Open POT Folder")
        self.pot_folder_btn.clicked.connect(self.open_report_folder)
        self.pot_folder_btn.setEnabled(False)

        top_bar.addWidget(back_btn)
        top_bar.addStretch()
        top_bar.addWidget(self.pot_folder_btn)
        layout.addLayout(top_bar)
        layout.addSpacing(4)

        # ===== Title & Subtitle =====
        title = QtWidgets.QLabel("Create POT")
        title.setObjectName("Title")
        layout.addWidget(title)
        layout.addSpacing(4)

        subtitle = QtWidgets.QLabel("Fill in the details below to generate a Proof of Test document")
        subtitle.setObjectName("Subtitle")
        subtitle.setWordWrap(True)
        layout.addWidget(subtitle)

        layout.addSpacing(24)

        # Shared field style so every input matches regardless of section
        field_style = f"""
            QLineEdit {{
                background-color: {tc['field_bg']};
                border: 1px solid {tc['field_border']};
                border-radius: 8px;
                padding: 9px 10px;
                color: {tc['field_text']};
                font-size: 13px;
            }}
            QLineEdit:focus {{
                border: 1px solid {tc['field_focus_border']};
            }}
            QDateEdit {{
                background-color: {tc['field_bg']};
                border: 1px solid {tc['field_border']};
                border-radius: 8px;
                padding: 9px 36px 9px 10px;
                color: {tc['field_text']};
                font-size: 13px;
            }}
            QDateEdit:focus {{
                border: 1px solid {tc['field_focus_border']};
            }}
            QDateEdit::drop-down {{
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 28px;
                border: none;
                border-left: 1px solid {tc['field_border']};
                border-top-right-radius: 8px;
                border-bottom-right-radius: 8px;
            }}
            QDateEdit::down-arrow {{
                width: 0;
                height: 0;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 6px solid {tc['field_label']};
            }}
        """
        section_label_style = f"""
            color: {tc['section_label']};
            font-size: 11px;
            font-weight: 700;
            letter-spacing: 1.5px;
        """
        field_label_style = f"color: {tc['field_label']}; font-size: 12px; font-weight: 600;"

        def make_field_block(label_text, widget):
            block = QtWidgets.QVBoxLayout()
            block.setSpacing(6)
            lbl = QtWidgets.QLabel(label_text)
            lbl.setStyleSheet(field_label_style)
            widget.setStyleSheet(field_style)
            block.addWidget(lbl)
            block.addWidget(widget)
            return block

        # ===== Section 1: Ticket Information =====
        ticket_section_label = QtWidgets.QLabel("TICKET INFORMATION")
        ticket_section_label.setStyleSheet(section_label_style)
        layout.addWidget(ticket_section_label)
        layout.addSpacing(12)

        self.pot_title_input = QtWidgets.QLineEdit()
        self.ticket_input = QtWidgets.QLineEdit()
        self.date_input = QtWidgets.QDateEdit()
        self.date_input.setCalendarPopup(True)
        self.date_input.setDate(QtCore.QDate.currentDate())
        self.module_input = QtWidgets.QLineEdit()
        self.url_input = QtWidgets.QLineEdit()

        ticket_grid = QtWidgets.QGridLayout()
        ticket_grid.setHorizontalSpacing(20)
        ticket_grid.setVerticalSpacing(16)
        ticket_grid.addLayout(make_field_block("POT Title", self.pot_title_input), 0, 0, 1, 2)
        ticket_grid.addLayout(make_field_block("Ticket No.", self.ticket_input), 1, 0)
        ticket_grid.addLayout(make_field_block("Date", self.date_input), 1, 1)
        ticket_grid.addLayout(make_field_block("Affected Module / Screens", self.module_input), 2, 0)
        ticket_grid.addLayout(make_field_block("URL", self.url_input), 2, 1)
        layout.addLayout(ticket_grid)

        layout.addSpacing(24)
        divider1 = QtWidgets.QFrame()
        divider1.setFixedHeight(1)
        divider1.setStyleSheet(f"background: {tc['divider']}; border: none;")
        layout.addWidget(divider1)
        layout.addSpacing(20)

        # ===== Section 2: Company Information =====
        company_section_label = QtWidgets.QLabel("COMPANY INFORMATION")
        company_section_label.setStyleSheet(section_label_style)
        layout.addWidget(company_section_label)
        layout.addSpacing(12)

        self.company_input = QtWidgets.QLineEdit()
        self.company_id_input = QtWidgets.QLineEdit()
        self.company_name_input = QtWidgets.QLineEdit()
        self.issue_list_input = QtWidgets.QLineEdit()

        company_grid = QtWidgets.QGridLayout()
        company_grid.setHorizontalSpacing(20)
        company_grid.setVerticalSpacing(16)
        company_grid.addLayout(make_field_block("Company", self.company_input), 0, 0)
        company_grid.addLayout(make_field_block("Company ID", self.company_id_input), 0, 1)
        company_grid.addLayout(make_field_block("Company Name", self.company_name_input), 1, 0)
        company_grid.addLayout(make_field_block("Issue List", self.issue_list_input), 1, 1)
        layout.addLayout(company_grid)

        layout.addSpacing(28)

        # ===== Buttons =====
        btn_layout = QtWidgets.QHBoxLayout()
        btn_layout.setSpacing(12)

        cancel_btn = QtWidgets.QPushButton("Cancel")
        cancel_btn.setFixedHeight(42)
        cancel_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        cancel_btn.clicked.connect(self.show_main_menu_page)
        cancel_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                border: 1.5px solid {tc['outline_border']};
                border-radius: 10px;
                padding: 10px 20px;
                font-size: 14px;
                font-weight: 600;
                color: {tc['outline_text']};
            }}
            QPushButton:hover {{
                border: 1.5px solid {tc['cancel_hover_border']};
                color: {tc['outline_hover_text']};
            }}
        """)

        self.generate_btn = QtWidgets.QPushButton("Generate POT")
        self.generate_btn.setFixedHeight(42)
        self.generate_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        self.generate_btn.clicked.connect(self.generate_pot)

        btn_layout.addStretch()
        btn_layout.addWidget(cancel_btn)
        btn_layout.addWidget(self.generate_btn)
        layout.addLayout(btn_layout)

        layout.addSpacing(16)

        # ===== Status label =====
        self.status_label = QtWidgets.QLabel("")
        self.status_label.setWordWrap(True)
        self.status_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignLeft)
        if is_windows_dark_mode():
            self.status_label.setStyleSheet("color: #4ade80; font-weight: 600; font-size: 12.5px;")
        else:
            self.status_label.setStyleSheet("color: #16a34a; font-weight: 600; font-size: 12.5px;")
        layout.addWidget(self.status_label)

        self.main_layout.addWidget(card, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

    def generate_pot(self):
        work_item_id = self.ticket_input.text().strip()

        if not work_item_id:
            if hasattr(self, "status_label"):
                self.status_label.setText("❌ Ticket No. is required to fetch TFS comment data.")
            return

        if not hasattr(self, "current_user"):
            if hasattr(self, "status_label"):
                self.status_label.setText("❌ No logged-in user data available.")
            return

        # Disable button + show feedback while fetching
        self.generate_btn.setEnabled(False)
        self.generate_btn.setText("Fetching ticket data...")
        if hasattr(self, "status_label"):
            self.status_label.setText("")
        QApplication.processEvents()

        self.pot_ticket_worker = PotTicketDataWorker(work_item_id, self.current_user)
        self.pot_ticket_worker.progress.connect(
            lambda msg: self.status_label.setText(msg) if hasattr(self, "status_label") else None
        )
        self.pot_ticket_worker.finished.connect(self._on_pot_ticket_data_fetched)
        self.pot_ticket_worker.error.connect(self._on_pot_ticket_data_error)
        self.pot_ticket_worker.start()


    def _on_pot_ticket_data_error(self, msg):
        self.generate_btn.setEnabled(True)
        self.generate_btn.setText("Generate POT")
        if hasattr(self, "status_label"):
            self.status_label.setText(msg)


    def _on_pot_ticket_data_fetched(self, raw_html):
        logger.info("POT ticket data received: %d chars of HTML", len(raw_html))
        self.generate_btn.setEnabled(True)
        self.generate_btn.setText("Generate POT")

        # --- Load the Word template ---
        pot_template_path = "POT_Template.docx"
        logger.info("Loading POT template: %s (cwd=%s)", pot_template_path, os.getcwd())
        if not os.path.isfile(pot_template_path):
            logger.error("POT template NOT FOUND: %s", pot_template_path)
            if hasattr(self, "status_label"):
                self.status_label.setText(f"❌ Template not found: {pot_template_path}")
            return
        try:
            doc = Document(pot_template_path)
            logger.info("POT template loaded (%d paragraph(s), %d table(s))", len(doc.paragraphs), len(doc.tables))
        except Exception:
            logger.exception("Failed to load POT template: %s", pot_template_path)
            if hasattr(self, "status_label"):
                self.status_label.setText(f"❌ Error loading template: {pot_template_path}")
            return

        # --- Get logged-in user's display name ---
        display_name = self.current_user.get("display_name", "") if hasattr(self, "current_user") else ""

        # --- Prepare text-only replacements dictionary ---
        replacements = {
            "{{POT_TITLE}}": self.pot_title_input.text(),
            "{{COMPANY}}": self.company_input.text(),
            "{{TICKET_NO}}": self.ticket_input.text(),
            "{{DATE}}": self.date_input.date().toString("MM-dd-yyyy"),
            "{{MODULE}}": self.module_input.text(),
            "{{URL}}": self.url_input.text(),
            "{{ISSUE_LIST}}": self.issue_list_input.text(),
            "{{COMPANY_ID}}": self.company_id_input.text(),
            "{{COMPANY_NAME}}": self.company_name_input.text(),
            "{{DISPLAYN}}": display_name,
        }
        logger.debug("POT replacements: %s", {k: v[:40] if isinstance(v, str) else v for k, v in replacements.items()})

        # --- Replace all plain-text placeholders everywhere (paragraphs + tables + text boxes) ---
        replace_placeholders_in_doc(doc, replacements)

        # --- Insert TICKETDATA (raw HTML, with images) into its paragraph ---
        self._insert_html_into_placeholder(doc, "{{TICKETDATA}}", raw_html)

        # --- Generate output file name safely ---
        ticket = self.ticket_input.text().strip().replace(" ", "_")
        company = self.company_input.text().strip().replace(" ", "_")
        output_filename = f"TFS{ticket}_WebPOS_{company}_POT_Internal_Passv1.0.docx"

        # Save into the same dated folder as status reports (MM-DD-YYYY)
        today_folder = datetime.now().strftime("%m-%d-%Y")
        os.makedirs(today_folder, exist_ok=True)
        output_path = os.path.join(today_folder, output_filename)

        try:
            doc.save(output_path)
            logger.info("POT saved successfully: %s", output_path)
        except Exception:
            logger.exception("Failed to save POT document: %s", output_path)
            if hasattr(self, "status_label"):
                self.status_label.setText(f"❌ Error saving POT: {output_path}")
            return

        # Track for the Open Folder button
        self.last_output_path = output_path
        if hasattr(self, "pot_folder_btn"):
            self.pot_folder_btn.setEnabled(True)

        if hasattr(self, "status_label"):
            self.status_label.setText(f"✅ POT generated and saved to {output_path}")
        else:
            print(f"POT generated: {output_path}")


    def _insert_html_into_placeholder(self, doc, placeholder, raw_html):
        """
        Finds the paragraph containing `placeholder` (in body paragraphs or
        table cells) and replaces it with rendered HTML content, including
        embedded images - reusing the same HTML->Word logic as ReportWorker.
        """
        logger.info("Looking for placeholder %r in document", placeholder)
        pat = self.current_user.get("pat", "") if hasattr(self, "current_user") else ""
        found = False

        def try_paragraph(paragraph):
            nonlocal found
            if placeholder in paragraph.text:
                logger.debug("Found %r in paragraph text", placeholder)
                for run in paragraph.runs:
                    run.text = ""
                add_html_content_to_paragraph(paragraph, raw_html, pat)
                found = True
                return True
            return False

        # Body paragraphs
        for para in doc.paragraphs:
            if try_paragraph(para):
                return

        # Table cells (including nested tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if try_paragraph(para):
                            return
                    for nested_table in cell.tables:
                        for nrow in nested_table.rows:
                            for ncell in nrow.cells:
                                for npara in ncell.paragraphs:
                                    if try_paragraph(npara):
                                        return

        if not found:
            logger.warning("Placeholder %r was NOT FOUND in any paragraph, table cell, or text box!", placeholder)

    def open_report_folder(self):
        if hasattr(self, "last_output_path"):
            folder = os.path.dirname(os.path.abspath(self.last_output_path))  # get absolute folder path
            print("Trying to open folder:", folder)  # debug print
            if os.path.exists(folder):
                os.startfile(folder)  # open folder in Explorer
            else:
                print("Folder does not exist:", folder)
        else:
            print("No last_output_path set yet")
    
    

    #CREATE REPORT PAGE
    def show_report_page(self):
        self.clear_layout()
        self.set_page_size(460, 700)

        card = QtWidgets.QWidget()
        card.setObjectName("Card")
        card.setFixedWidth(400)
        layout = QtWidgets.QVBoxLayout(card)
        layout.setContentsMargins(28, 28, 28, 28)
        layout.setSpacing(12)


        tc = self._theme_colors()

        #FOLDER LAYOUT
        top_layout = QtWidgets.QHBoxLayout()

        back_btn = QtWidgets.QToolButton()
        back_btn.setText("←")
        back_btn.setToolTip("Back to Main Menu")
        back_btn.setCursor(QtCore.Qt.CursorShape.PointingHandCursor)
        back_btn.clicked.connect(self.show_main_menu_page)
        back_btn.setStyleSheet(f"""
            QToolButton {{
                border: none;
                font-size: 18px;
                font-weight: bold;
                color: {tc['back_color']};
            }}
            QToolButton:hover {{
                color: {tc['back_hover']};
            }}
        """)

        title = QtWidgets.QLabel("Generate Status Report")
        title.setObjectName("Title")

        self.open_folder_btn = QtWidgets.QToolButton()
        self.open_folder_btn.setIcon(self.style().standardIcon(QtWidgets.QStyle.StandardPixmap.SP_DirOpenIcon))
        self.open_folder_btn.setToolTip("Open Report Folder")
        self.open_folder_btn.clicked.connect(self.open_report_folder)

        top_layout.addWidget(back_btn)
        top_layout.addWidget(title)
        top_layout.addStretch()
        top_layout.addWidget(self.open_folder_btn)

        layout.addLayout(top_layout)
        # Subtitle
        subtitle = QtWidgets.QLabel("Fill in the details below to create a report")
        subtitle.setObjectName("Subtitle")
        subtitle.setWordWrap(True)
        layout.addWidget(subtitle)

        # Work Item ID
        self.work_item_input = QtWidgets.QLineEdit()
        self.work_item_input.setPlaceholderText("Work Item ID")
        layout.addWidget(self.work_item_input)

        # Client Name
        self.client_input = QtWidgets.QLineEdit()
        self.client_input.setPlaceholderText("Client Name")
        layout.addWidget(self.client_input)

        # Location type
        loc_layout = QtWidgets.QHBoxLayout()
        loc_label = QtWidgets.QLabel("Location Type:")
        loc_label.setStyleSheet(f"color: {tc['field_label']};")
        loc_layout.addWidget(loc_label)

        self.onsite_radio = QtWidgets.QRadioButton("Onsite")
        self.offsite_radio = QtWidgets.QRadioButton("Offsite / Remote")
        self.offsite_radio.setChecked(True)
        loc_layout.addWidget(self.onsite_radio)
        loc_layout.addWidget(self.offsite_radio)
        layout.addLayout(loc_layout)

        # Generate Button
        self.generate_report_btn = QtWidgets.QPushButton("Generate Report")
        self.generate_report_btn.clicked.connect(self.generate_report)
        layout.addWidget(self.generate_report_btn)

        for field in card.findChildren(QtWidgets.QLineEdit):
            field.returnPressed.connect(self.generate_report_btn.click)

        # Output / status area
        self.output_area = QtWidgets.QTextEdit()
        self.output_area.setReadOnly(True)
        self.output_area.setStyleSheet(f"background: {tc['output_bg']}; color: {tc['output_text']};")
        layout.addWidget(self.output_area)

        self.status_label = QLabel("")  # initially empty
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignLeft)
        if is_windows_dark_mode():
            self.status_label.setStyleSheet("color: #4ade80; font-weight: bold;")
        else:
            self.status_label.setStyleSheet("color: #16a34a; font-weight: bold;")
        layout.addWidget(self.status_label)  # add it to the bottom

        # Create the button but keep it hidden initially
        self.create_azure_btn = QtWidgets.QPushButton("Create Azure Status Report")
        self.create_azure_btn.setVisible(False)  # hidden until DOCX is saved
        self.create_azure_btn.clicked.connect(self.on_create_azure_btn_clicked)
        layout.addWidget(self.create_azure_btn, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

        self.main_layout.addWidget(card, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

    def on_create_azure_btn_clicked(self):
        # Make sure we have report data
        if not hasattr(self, "last_report_data"):
            self.output_area.append("❌ No report data available to create PBI.")
            return
        user_data = self.current_user
        pat = user_data["pat"]

        # Build a default report title
        report_title = f"Ticket No. {self.last_report_data['work_item_id']} - Status Report for {self.last_report_data['client_name']}"

        # Open the PBI input dialog
        dialog = PBIFormDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            values = dialog.get_values()  # get user input from the dialog

            # Call your PBI creation method with GUI values
            self.create_pbi_interactively_tfs(
                pat,
                report_title,
                self.last_report_data["root_cause"],
                self.last_report_data["preventive_action"],
                self.last_report_data["next_step"],
                self.last_report_data["raw_html"],
                back_job=values["back_job"],
                billing_status=values["billing_status"],
                product_type=values["product_type"],
                support_type=values["support_type"],
                confirmed_by=values["confirmed_by"]
            )
    


    def create_pbi_interactively_tfs(
            self,
            pat,
            report_title,
            root_cause,
            preventive_action,
            next_step,
            raw_html,
            back_job,
            billing_status,
            product_type,
            support_type,
            confirmed_by,
            timezone_str="Asia/Kuala_Lumpur",
        ):
        """
        Create a Product Backlog Item (PBI) in TFS (Azure DevOps Server)
        using GUI-provided values and show result in a popup.
        """
        logger.info("=== Creating PBI: %r ===", report_title)

        # --- Static config ---
        collection_url = "https://tfs.alliancewebpos.com/tfs/WebPOSCollection"
        project = "WebPOS_EES"

        user_data = self.current_user
        assigned_to = user_data["display_name"]

        # --- Get current time ---
        try:
            tz = pytz.timezone(timezone_str)
            now = datetime.now(tz)
            current_time = now.strftime("%m/%d/%Y %I:%M %p %Z").lstrip("0").replace("/0", "/")
        except Exception as e:
            now = datetime.now()
            current_time = now.strftime("%m/%d/%Y %I:%M %p")
            logger.warning("Timezone '%s' not found, using local time: %s", timezone_str, e)

        # --- REST API URL ---
        url = f"{collection_url}/{project}/_apis/wit/workitems/$Product%20Backlog%20Item?api-version=5.1"
        headers = {'Content-Type': 'application/json-patch+json'}

        # --- Build request body ---
        body = [
            {"op": "add", "path": "/fields/System.Title", "value": report_title},
            {"op": "add", "path": "/fields/System.Description", "value": raw_html},
            {"op": "add", "path": "/fields/System.AssignedTo", "value": assigned_to},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReportBillingStatus", "value": billing_status},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReportProducttype", "value": product_type},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReportSupportType", "value": support_type},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReportConfirmedBy", "value": confirmed_by},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReport_RootCause", "value": root_cause or "N/A"},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReport_PreventiveAction", "value": preventive_action or "N/A"},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReport_NextSteps", "value": next_step or "N/A"},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReportBackjob", "value": back_job},
            {"op": "add", "path": "/fields/Alliance.Groundup.Workitem.StatusReportDate", "value": current_time}
        ]
        logger.debug("PBI body fields: billing=%r, product=%r, support=%r, confirmed_by=%r, back_job=%r",
                      billing_status, product_type, support_type, confirmed_by, back_job)

        # --- Make API call ---
        try:
            logger.debug("POST %s", url)
            response = requests.post(
                url,
                auth=HTTPBasicAuth('', pat),
                headers=headers,
                json=body
            )

            if response.status_code in (200, 201):
                data = response.json()

                pbi_id = data.get("id")

                # Prefer human-friendly web URL
                pbi_url = (
                    data.get("_links", {})
                        .get("html", {})
                        .get("href")
                    or data.get("url")
                )

                logger.info("PBI created successfully: id=%s, url=%s", pbi_id, pbi_url)

                dialog = PBICreatedDialog(pbi_id, pbi_url, self)
                dialog.exec()

                self.reset_generate_report_page()

            else:
                logger.error("PBI creation failed: status=%d, body=%s", response.status_code, response.text[:500])
                AppDialog(
                    "PBI Creation Failed",
                    f"Status Code: {response.status_code}\n\n{response.text[:300]}",
                    self
                ).exec()

        except Exception as e:
            logger.exception("Error creating PBI")
            AppDialog("PBI Creation Error", f"Error creating PBI: {e}", self).exec()

    
    def reset_generate_report_page(self):
        # Clear text fields
        self.work_item_input.clear()
        self.client_input.clear()

        # Reset radio buttons
        self.onsite_radio.setChecked(True)
        self.offsite_radio.setChecked(False)

        # Clear output area
        self.output_area.clear()

        # Hide Azure button again
        if hasattr(self, "create_azure_btn"):
            self.create_azure_btn.setVisible(False)
            self.create_azure_btn.setEnabled(False)

        # Clear saved report data
        if hasattr(self, "last_report_data"):
            del self.last_report_data

        # Optional status message
        self.status_label.setText("")


    def generate_report(self):
        work_item_id = self.work_item_input.text().strip()
        client_name = self.client_input.text().strip()
        location_type1 = "X" if self.onsite_radio.isChecked() else ""
        location_type2 = "X" if self.offsite_radio.isChecked() else ""

        if not work_item_id or not client_name:
            self.output_area.append("Work Item ID and Client Name cannot be empty.")
            return

        # Disable inputs to prevent duplicate submissions
        self.generate_report_btn.setEnabled(False)
        self.generate_report_btn.setText("Generating...")
        self.work_item_input.setEnabled(False)
        self.client_input.setEnabled(False)
        self.onsite_radio.setEnabled(False)
        self.offsite_radio.setEnabled(False)

        self.output_area.append("Starting report generation...")

        # Create worker
        self.worker = ReportWorker(
            work_item_id, client_name, location_type1, location_type2, self.current_user, template_path
        )

        # Connect signals
        self.worker.progress.connect(lambda msg: self.output_area.append(msg))
        self.worker.finished.connect(self.handle_report_finished)
        self.worker.report_saved.connect(self.on_report_saved)

        # Start worker
        self.worker.start()

    def on_report_saved(self, path):
        self.last_output_path = path
        self.open_folder_btn.setEnabled(True)
        print("Report saved at:", path)


 
 #detect windows theme
def is_windows_dark_mode():
    try:
        registry = winreg.ConnectRegistry(None, winreg.HKEY_CURRENT_USER)
        key = winreg.OpenKey(
            registry,
            r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize"
        )

        # 0 = Dark Mode
        # 1 = Light Mode
        value, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")

        return value == 0
    except Exception:
        return True

# ------------------- Splash Screen -------------------
class SplashScreen(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint)
        self.setFixedSize(300, 150)
        if is_windows_dark_mode():
            self.setStyleSheet("""
                background-color: #111a22;
                color: #e6eef5;
                font-family: Inter, Segoe UI, Arial;
                font-size: 14px;
            """)
        else:
            self.setStyleSheet("""
                background-color: #f5f7fa;
                color: #111827;
                font-family: Inter, Segoe UI, Arial;
                font-size: 14px;
            """)
        layout = QVBoxLayout()
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label = QLabel("Loading application...")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(label)
        self.setLayout(layout)    



if __name__ == "__main__":
    import sys
    from PyQt6.QtWidgets import QApplication

    # Ensure relative paths resolve from the exe's directory, not the
    # current working directory (matters when launching via shortcut).
    if getattr(sys, 'frozen', False):
        os.chdir(os.path.dirname(sys.executable))

    logger.info("=== Application starting (cwd=%s) ===", os.getcwd())
    logger.info("Python %s, frozen=%s", sys.version, getattr(sys, 'frozen', False))

    app = QApplication(sys.argv)
    app.setWindowIcon(QtGui.QIcon("asset/death.ico"))

    # Show splash screen first
    splash = SplashScreen()
    splash.show()

    # Use QTimer to simulate loading
    def load_main_window():
        logger.info("Loading main window")
        # Here you can import heavy modules if needed (e.g., pandas)
        window = LoginWindow()
        window.show()
        splash.close()
        app.main_window = window  # keep reference
        logger.info("Main window loaded and shown")

        # Check for updates after UI is visible (non-blocking)
        def check_update():
            update = _check_for_update()
            if update:
                logger.info("Update available: %s", update["version"])
                window._update_info = update
                dlg = UpdateDialog(update["version"], update["body"], parent=window)
                dlg.exec()
            else:
                logger.info("No update available")

        QTimer.singleShot(2000, check_update)

    # Start loading after 100ms
    QTimer.singleShot(100, load_main_window)

    sys.exit(app.exec())
